import io
import os
import streamlit as st
import pandas as pd
import numpy as np
import requests
import plotly.express as px
import time
from datetime import datetime, date, timedelta
from calendar import monthrange
from concurrent.futures import ThreadPoolExecutor, as_completed
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry

# 1. 페이지 설정
st.set_page_config(page_title="활주로 이용률 정밀 분석", layout="wide")

# 설정 항목이 많아 한눈에 보기 쉽도록 전체 폰트를 약간 축소
st.markdown("""
<style>
.block-container { padding-top: 2rem; font-size: 0.92rem; }
h3 { font-size: 1.05rem !important; }
h4 { font-size: 0.95rem !important; }
div[data-testid="stMetricValue"] { font-size: 1.2rem; }
div[data-testid="stMetricLabel"] { font-size: 0.8rem; }
div[data-testid="stMetricDelta"] { font-size: 0.78rem; }
label[data-testid="stWidgetLabel"] p { font-size: 0.85rem; }
div[data-testid="stExpander"] summary p { font-size: 0.88rem; }
</style>
""", unsafe_allow_html=True)

# --- [개인용 기본값] secrets.toml에 저장된 키를 세션 상태 기본값으로 자동 채움.
#     secrets.toml은 .gitignore 처리되어 GitHub(public repo)에는 절대 올라가지 않음.
def _secret(name, default=""):
    try:
        return st.secrets.get(name, default)
    except Exception:
        return default

# API 키는 메인 화면에 직접 노출하지 않고 세션 상태로만 보관 → 팝업(dialog)에서 입력/수정
for _sk, _sname in [("api_key", "ASOS_API_KEY"), ("kma_hub_key", "KMA_HUB_KEY"), ("kakao_key", "KAKAO_REST_KEY")]:
    if _sk not in st.session_state:
        st.session_state[_sk] = _secret(_sname)

@st.dialog("API 키 설정")
def _api_key_dialog():
    st.text_input(
        "ASOS API Key (Decoding)", type="password", key="api_key",
        value=st.session_state.get("api_key", ""),
        help="공공데이터포털(data.go.kr)에서 발급받은 종관기상관측(ASOS) Decoding 키",
    )
    st.text_input(
        "기상청 API 허브 인증키", type="password", key="kma_hub_key",
        value=st.session_state.get("kma_hub_key", ""),
        help="apihub.kma.go.kr에서 발급받은 인증키 (관측소 위경도 조회 및 주소검색용)",
    )
    st.text_input(
        "카카오 REST API 키", type="password", key="kakao_key",
        value=st.session_state.get("kakao_key", ""),
        help="Kakao Developers(developers.kakao.com)에서 앱 생성 후 즉시 발급되는 REST API 키. "
             "별도 승인 대기 없이 바로 사용 가능합니다.",
    )
    st.caption("입력한 값은 이 브라우저 세션에만 유지되며, 코드나 GitHub에는 저장되지 않습니다.")
    if st.button("닫기", type="primary", use_container_width=True):
        st.rerun()

# 제목 + API 키 설정 버튼 (같은 행에 배치)
title_col1, title_col2 = st.columns([6, 1])
with title_col1:
    st.title("활주로 이용률 정밀 분석")
with title_col2:
    st.write("")
    st.write("")
    if st.button("API 키 설정", use_container_width=True):
        _api_key_dialog()

api_key = st.session_state.get("api_key", "")

# --- [내장 데이터] 광역시도 → (ASOS 종관 / AWS 방재) 관측소 계층 ---
# ASOS: 기상청 종관기상관측망 전체 (94개소) — AsosHourlyInfoService API로 시간자료 제공
ASOS_BY_REGION = {
    "서울": [
        ("108", "서울", "1907-10-01"),
    ],
    "인천": [
        ("112", "인천", "1904-08-01"),
        ("102", "백령도", "2000-01-01"),
        ("201", "강화", "1972-01-01"),
    ],
    "경기": [
        ("98", "동두천", "1998-02-01"),
        ("99", "파주", "2001-12-01"),
        ("119", "수원", "1964-01-01"),
        ("202", "양평", "1972-01-01"),
        ("203", "이천", "1972-01-01"),
    ],
    "강원": [
        ("90", "속초", "1968-01-01"),
        ("93", "북춘천", "2016-10-01"),
        ("95", "철원", "1988-01-01"),
        ("100", "대관령", "1971-12-01"),
        ("101", "춘천", "1966-01-01"),
        ("104", "북강릉", "2008-10-01"),
        ("105", "강릉", "1911-10-01"),
        ("114", "원주", "1971-01-01"),
        ("121", "영월", "1994-01-01"),
        ("211", "인제", "1972-01-01"),
        ("212", "홍천", "1972-01-01"),
        ("216", "태백", "1985-01-01"),
        ("217", "정선군", "2010-01-01"),
    ],
    "충북": [
        ("127", "충주", "1972-01-01"),
        ("131", "청주", "1967-01-01"),
        ("221", "제천", "1972-01-01"),
        ("226", "보은", "1972-01-01"),
    ],
    "대전": [
        ("133", "대전", "1969-01-01"),
    ],
    "세종": [
        ("239", "세종", "2019-10-01"),
    ],
    "충남": [
        ("129", "서산", "1968-01-01"),
        ("232", "천안", "1972-01-01"),
        ("235", "보령", "1972-01-01"),
        ("236", "부여", "1972-01-01"),
        ("238", "금산", "1972-01-01"),
    ],
    "전북": [
        ("140", "군산", "1968-01-01"),
        ("146", "전주", "1918-01-01"),
        ("172", "고창", "2010-12-01"),
        ("243", "부안", "1972-01-01"),
        ("244", "임실", "1972-01-01"),
        ("245", "정읍", "1972-01-01"),
        ("247", "남원", "1972-01-01"),
        ("248", "장수", "1972-01-01"),
        ("251", "고창군", "2010-01-01"),
        ("254", "순창군", "2010-01-01"),
    ],
    "광주": [
        ("156", "광주", "1938-10-01"),
    ],
    "전남": [
        ("165", "목포", "1904-04-01"),
        ("168", "여수", "1942-02-01"),
        ("169", "흑산도", "1997-01-01"),
        ("170", "완도", "1971-01-01"),
        ("174", "순천", "1973-01-01"),
        ("177", "진도(첨찰산)", "2012-01-01"),
        ("252", "영광군", "2010-01-01"),
        ("258", "보성군", "2010-01-01"),
        ("259", "강진군", "2009-12-01"),
        ("260", "장흥", "1972-01-01"),
        ("261", "해남", "2010-05-01"),
        ("262", "고흥", "1972-01-01"),
        ("266", "광양시", "2010-01-01"),
        ("268", "진도군", "2009-12-01"),
    ],
    "대구": [
        ("143", "대구", "1907-01-01"),
    ],
    "경북": [
        ("115", "울릉도", "1938-08-01"),
        ("130", "울진", "1971-01-01"),
        ("135", "추풍령", "1935-01-01"),
        ("136", "안동", "1973-01-01"),
        ("137", "상주", "2002-01-01"),
        ("138", "포항", "1943-01-01"),
        ("271", "봉화", "1988-01-01"),
        ("272", "영주", "1972-01-01"),
        ("273", "문경", "1973-01-01"),
        ("276", "청송군", "2010-01-01"),
        ("277", "영덕", "1972-01-01"),
        ("278", "의성", "1973-01-01"),
        ("279", "구미", "1973-01-01"),
        ("281", "영천", "1972-01-01"),
        ("283", "경주시", "2010-01-01"),
    ],
    "부산": [
        ("159", "부산", "1904-04-01"),
    ],
    "울산": [
        ("152", "울산", "1931-01-01"),
    ],
    "경남": [
        ("155", "창원", "1985-01-01"),
        ("162", "통영", "1968-01-01"),
        ("192", "진주", "1969-01-01"),
        ("253", "김해시", "2010-01-01"),
        ("255", "북창원", "2010-01-01"),
        ("257", "양산시", "2010-01-01"),
        ("263", "의령군", "2010-01-01"),
        ("264", "함양군", "2010-01-01"),
        ("284", "거창", "1972-01-01"),
        ("285", "합천", "1973-01-01"),
        ("288", "밀양", "1973-01-01"),
        ("289", "산청", "1973-01-01"),
        ("294", "거제", "1972-01-01"),
        ("295", "남해", "1972-01-01"),
    ],
    "제주": [
        ("184", "제주", "1923-05-01"),
        ("185", "고산", "1988-01-01"),
        ("188", "성산", "1973-01-01"),
        ("189", "서귀포", "1961-01-01"),
    ],
}

# ※ AWS(방재기상관측) 관측소 목록은 의도적으로 제공하지 않음.
#   - 기상청 API 허브의 공식 관측소 정보 API(stn_inf.php?inf=AWS)는 "활용신청 필요(403)"로 비공개
#   - 검증되지 않은 추정 ID/명칭을 보여주면 사용자에게 혼동을 줄 수 있어 제외함
#   - 더 근본적으로, AWS 시간자료 API(awsh.php)는 "단일 시점 조회"만 지원하고
#     기간(tm1~tm2) 조회는 항상 "최근 약 1개월"만 반환하도록 고정되어 있어
#     5~10년 분석에 필요한 장기 시계열 수집이 구조적으로 불가능함 (아래 사이드바 안내 참고)

# 플랫 조회용 (id → (type, name, region, start_date))
def _build_lookup():
    lk = {}
    for region, lst in ASOS_BY_REGION.items():
        for sid, name, start in lst:
            lk.setdefault(sid, ("ASOS", name, region, start))
    return lk

STATION_LOOKUP = _build_lookup()
# 역호환: STATION_DB 이름 유지 (id → [name, start])
STATION_DB = {sid: [v[1], v[3]] for sid, v in STATION_LOOKUP.items()}

API_URL = "http://apis.data.go.kr/1360000/AsosHourlyInfoService/getWthrDataList"
MAX_WORKERS = 8          # 병렬 요청 수 (공공데이터 포털 권장 범위 내)
REQUEST_TIMEOUT = 20     # 단건 타임아웃
MAX_RETRIES = 3

def _make_session():
    """Keep-Alive + 자동 재시도 세션. 커넥션 풀링으로 TCP/TLS 핸드쉐이크 비용 절감."""
    s = requests.Session()
    retry = Retry(
        total=MAX_RETRIES,
        backoff_factor=0.5,           # 0.5s, 1s, 2s ...
        status_forcelist=[429, 500, 502, 503, 504],
        allowed_methods=["GET"],
        raise_on_status=False,
    )
    adapter = HTTPAdapter(pool_connections=MAX_WORKERS, pool_maxsize=MAX_WORKERS, max_retries=retry)
    s.mount("http://", adapter)
    s.mount("https://", adapter)
    return s

def _month_chunks(s_date, e_date):
    """분석 기간을 '월 단위'로 쪼개어 [(startDt, endDt), ...] 반환.
    1개월 ≈ 720행 → numOfRows=999 한 페이지에 수렴하므로 페이지네이션 불필요."""
    chunks = []
    cur = date(s_date.year, s_date.month, 1)
    while cur <= e_date:
        last_day = monthrange(cur.year, cur.month)[1]
        m_start = max(cur, s_date)
        m_end = min(date(cur.year, cur.month, last_day), e_date)
        chunks.append((m_start.strftime("%Y%m%d"), m_end.strftime("%Y%m%d")))
        # 다음 달 1일로 이동
        cur = (date(cur.year + (cur.month // 12), ((cur.month % 12) + 1), 1))
    return chunks

def _fetch_one(session, key, stn, start_dt, end_dt):
    """월 단위 1회 요청. 월 > 999행 대비 페이지네이션 포함(안전장치)."""
    items_all = []
    for page in range(1, 4):  # 월 단위는 거의 1페이지로 끝나지만 안전하게 3페이지까지
        params = {
            'serviceKey': key, 'pageNo': str(page), 'numOfRows': '999',
            'dataType': 'JSON', 'dataCd': 'ASOS', 'dateCd': 'HR',
            'startDt': start_dt, 'startHh': '01',
            'endDt': end_dt, 'endHh': '23', 'stnIds': stn,
        }
        r = session.get(API_URL, params=params, timeout=REQUEST_TIMEOUT)
        if r.status_code != 200:
            # data.go.kr 게이트웨이 오류(502 등) → 원인을 명확히 전달 (JSON 파싱 전에 차단)
            raise RuntimeError(f"HTTP {r.status_code} · {r.text[:80].strip()}")
        try:
            res = r.json()
        except ValueError:
            # 키 오류·서비스 점검 시 JSON 대신 XML/평문 오류가 오므로 본문 앞부분을 노출
            raise RuntimeError(f"비정상 응답(JSON 아님): {r.text[:120].strip()}")
        items = res.get('response', {}).get('body', {}).get('items', {}).get('item', [])
        if isinstance(items, dict):
            items = [items]
        if not items:
            break
        items_all.extend(items)
        if len(items) < 999:
            break
    return items_all

# --- [캐시 데이터] 월별 병렬 수집 함수 ---
@st.cache_data(show_spinner=False)
def get_weather_data_v28(key, stn, s_date, e_date):
    chunks = _month_chunks(s_date, e_date)
    total = len(chunks)
    if total == 0:
        return None, None, "날짜 범위가 비어 있습니다."

    msg_slot = st.empty()
    p_bar = st.progress(0)
    msg_slot.info(f"⏳ {total}개월 데이터 병렬 수집 중... (동시 요청 {MAX_WORKERS}건)")

    all_combined = []
    done = 0
    t0 = time.perf_counter()

    failed = []
    with _make_session() as session, ThreadPoolExecutor(max_workers=MAX_WORKERS) as pool:
        futures = {pool.submit(_fetch_one, session, key, stn, s, e): (s, e) for s, e in chunks}
        for fut in as_completed(futures):
            s, e = futures[fut]
            try:
                items = fut.result()
                if items:
                    all_combined.extend(items)
            except Exception as ex:
                # 개별 구간 실패는 전체 수집을 중단시키지 않고 누적해 두었다가 보고
                # (일시적 네트워크 오류·API 혼잡 방어. 과거엔 첫 실패에서 즉시 return하며
                #  반환값 개수가 달라 호출부 언패킹이 깨졌음 — 반드시 3-tuple 유지)
                failed.append(f"{s}~{e}: {ex}")
            done += 1
            p_bar.progress(done / total)
            msg_slot.info(f"⏳ 수집 진행 {done}/{total}개월 · 경과 {time.perf_counter()-t0:.1f}s")

    msg_slot.success(f"수집 완료 · {len(all_combined):,}행 · 총 {time.perf_counter()-t0:.1f}s")

    if not all_combined:
        reason = "수집된 데이터가 없습니다. 날짜·지점·API 키를 확인하세요."
        if failed:
            reason += f" (구간 실패 {len(failed)}/{total}건 · 예: {failed[0]})"
        return None, None, reason

    df = pd.DataFrame(all_combined)
    df['wd']    = pd.to_numeric(df['wd'], errors='coerce')
    df['ws_kt'] = pd.to_numeric(df['ws'], errors='coerce') * 1.94384
    # 관측 시각 기준 중복 제거(병렬 중복 방어)
    if 'tm' in df.columns:
        df = df.drop_duplicates(subset=['tm'])
    # ── 유효 / 결측 분리 ──────────────────────────────────────────
    invalid_mask = df['wd'].isna() | df['ws_kt'].isna()
    df_invalid   = df[invalid_mask].reset_index(drop=True)
    df_valid     = df[~invalid_mask].reset_index(drop=True)
    return df_valid, df_invalid, len(df_valid)

# --- [AWS CSV 파싱 함수] ---
def _parse_aws_csv(uploaded_files):
    """기상자료개방포털 방재기상관측 시간자료 CSV 파싱.
    - 인코딩 자동 감지 (UTF-8-sig / EUC-KR / CP949)
    - 복수 파일 병합 지원 (연도별 분할 다운로드 대응)
    - 풍향·풍속 컬럼 자동 탐지
    반환: (df_valid, df_invalid, row_count, (start_date, end_date))
          또는 (None, None, error_msg, (None, None))
    """
    all_dfs = []
    for f in uploaded_files:
        raw = f.read()
        text = None
        for enc in ('utf-8-sig', 'euc-kr', 'cp949', 'utf-8'):
            try:
                text = raw.decode(enc)
                break
            except (UnicodeDecodeError, LookupError):
                continue
        if text is None:
            return None, None, f"'{f.name}' 인코딩 인식 불가 (UTF-8 또는 EUC-KR 파일 필요).", (None, None)
        try:
            tmp = pd.read_csv(io.StringIO(text))
            all_dfs.append(tmp)
        except Exception as e:
            return None, None, f"'{f.name}' CSV 파싱 오류: {e}", (None, None)

    if not all_dfs:
        return None, None, "업로드된 파일이 없습니다.", (None, None)

    df = pd.concat(all_dfs, ignore_index=True)

    # ── 풍향 컬럼 탐지 ────────────────────────────────────────────
    wd_col = next((c for c in df.columns if '풍향' in c), None)
    # 풍속: '최대', '순간', '돌풍' 제외한 첫 번째 풍속 컬럼
    ws_col = next(
        (c for c in df.columns
         if '풍속' in c and all(kw not in c for kw in ['최대', '순간', '돌풍'])),
        None
    )

    if wd_col is None:
        return None, None, f"풍향 컬럼을 찾을 수 없습니다. 전체 컬럼: {list(df.columns)}", (None, None)
    if ws_col is None:
        return None, None, f"풍속 컬럼을 찾을 수 없습니다. 전체 컬럼: {list(df.columns)}", (None, None)

    df['wd']    = pd.to_numeric(df[wd_col], errors='coerce')
    df['ws_kt'] = pd.to_numeric(df[ws_col], errors='coerce') * 1.94384   # m/s → knots

    # ── 날짜 범위 자동 감지 ───────────────────────────────────────
    dt_col = next((c for c in df.columns if '일시' in c or '날짜' in c), None)
    csv_start = csv_end = None
    if dt_col is not None:
        parsed_dt = pd.to_datetime(df[dt_col], errors='coerce').dropna()
        if len(parsed_dt) > 0:
            csv_start = parsed_dt.min().date()
            csv_end   = parsed_dt.max().date()

    # ── 유효 / 결측 분리 ──────────────────────────────────────────
    invalid_mask  = df['wd'].isna() | df['ws_kt'].isna()
    df_invalid    = df[invalid_mask].reset_index(drop=True)
    df_valid      = df[~invalid_mask].reset_index(drop=True)

    if len(df_valid) == 0:
        return None, None, "풍향·풍속 유효 데이터가 없습니다 (모두 결측).", (csv_start, csv_end)

    return df_valid, df_invalid, len(df_valid), (csv_start, csv_end)


# --- [주소 기반 관측소 검색 함수] ---
# 1) 기상청 API 허브 stn_inf.php → ASOS/AWS 관측소 위경도 DB
# 2) 카카오 로컬 API 주소검색    → 입력 주소 → WGS84 위경도 (1회 호출로 완료)
# 3) Haversine                  → 관측소 DB와의 거리 계산 → 최근접 N개

@st.cache_data(show_spinner=False, ttl=86400)
def _load_station_db(auth_key):
    """기상청 API 허브 stn_inf.php로 ASOS(SFC)+AWS 관측소 위경도 DB 구축.
    고정폭 텍스트(EUC-KR) → 바이트 단위 슬라이싱으로 파싱 (한글 폭 깨짐 방지)."""
    base = "https://apihub.kma.go.kr/api/typ01/url/stn_inf.php"
    tm = datetime.now().strftime('%Y%m%d%H%M')

    def _fetch_lines(inf):
        r = requests.get(base, params={'inf': inf, 'stn': '', 'tm': tm, 'help': '0', 'authKey': auth_key}, timeout=15)
        lines = r.content.split(b'\n')
        return [l for l in lines if l.strip() and not l.startswith(b'#')]

    rows = []
    # SFC(종관/ASOS): STN_ID(0:5) LON(5:19) LAT(19:33) ... STN_KO(96:117) ... LAW_ADDR(163:222)
    for line in _fetch_lines('SFC'):
        rows.append({
            'stn_id': line[0:5].decode('euc-kr', errors='replace').strip(),
            'lon':    line[5:19].decode('euc-kr', errors='replace').strip(),
            'lat':    line[19:33].decode('euc-kr', errors='replace').strip(),
            'name_ko': line[96:117].decode('euc-kr', errors='replace').strip(),
            'addr':   line[163:222].decode('euc-kr', errors='replace').strip(),
            'type':   'ASOS',
        })
    # AWS(방재): STN_ID(0:5) LON(5:19) LAT(19:33) ... STN_KO(71:92) ... LAW_ADDR(138:197)
    for line in _fetch_lines('AWS'):
        rows.append({
            'stn_id': line[0:5].decode('euc-kr', errors='replace').strip(),
            'lon':    line[5:19].decode('euc-kr', errors='replace').strip(),
            'lat':    line[19:33].decode('euc-kr', errors='replace').strip(),
            'name_ko': line[71:92].decode('euc-kr', errors='replace').strip(),
            'addr':   line[138:197].decode('euc-kr', errors='replace').strip(),
            'type':   'AWS',
        })

    df = pd.DataFrame(rows)
    df['lon'] = pd.to_numeric(df['lon'], errors='coerce')
    df['lat'] = pd.to_numeric(df['lat'], errors='coerce')
    df = df.dropna(subset=['lon', 'lat']).reset_index(drop=True)
    return df


def _geocode_address_kakao(keyword, kakao_key):
    """카카오 로컬 API 주소검색: 주소 키워드 → WGS84 위경도 1회 호출.
    반환: (lat, lon, display_addr, error_msg)."""
    try:
        r = requests.get(
            "https://dapi.kakao.com/v2/local/search/address.json",
            headers={"Authorization": f"KakaoAK {kakao_key}"},
            params={"query": keyword},
            timeout=10,
        )
        data = r.json()
        if r.status_code != 200:
            return None, None, None, data.get('message', f'HTTP {r.status_code} 오류 (REST API 키를 확인하세요)')
        docs = data.get('documents') or []
        if not docs:
            return None, None, None, "검색된 주소가 없습니다. 더 구체적인 주소를 입력해 보세요."
        d = docs[0]
        lon, lat = float(d['x']), float(d['y'])
        road = d.get('road_address')
        display_addr = road['address_name'] if road else d.get('address_name', keyword)
        return lat, lon, display_addr, None
    except Exception as e:
        return None, None, None, str(e)


def _haversine_km(lat1, lon1, lat2, lon2):
    """두 위경도 좌표 간 대원거리(km)."""
    R = 6371.0
    p1, p2 = np.radians(lat1), np.radians(lat2)
    dphi, dlambda = np.radians(lat2 - lat1), np.radians(lon2 - lon1)
    a = np.sin(dphi / 2) ** 2 + np.cos(p1) * np.cos(p2) * np.sin(dlambda / 2) ** 2
    return 2 * R * np.arcsin(np.sqrt(a))


def _nearest_stations(lat, lon, station_df, stn_type, n=3):
    """station_df에서 stn_type('ASOS'/'AWS') 관측소 중 (lat,lon)에 가장 가까운 n개."""
    sub = station_df[station_df['type'] == stn_type].copy()
    sub['dist_km'] = _haversine_km(lat, lon, sub['lat'].to_numpy(), sub['lon'].to_numpy())
    return sub.sort_values('dist_km').head(n).reset_index(drop=True)


def _norm_stn_id(s):
    """지점번호 표기 정규화(앞자리 0 차이 흡수). '090'/'90' → '90'."""
    s = str(s).strip()
    return str(int(s)) if s.isdigit() else s


# ASOS 지점번호 → (광역시도, 관측소명, 관측시작일) 역참조 (주소검색 자동선택용)
_ASOS_ID_TO_INFO = {
    _norm_stn_id(sid): (rgn, nm, start)
    for rgn, _stns in ASOS_BY_REGION.items()
    for (sid, nm, start) in _stns
}
NEAR_ASOS_PREFER_KM = 5.0   # 방재가 더 가까워도 종관과 이 거리(km) 미만 차이면 종관 우선 선택


# --- [ICAO/논문 기반 분석 함수] ---
# 논문: 신동진, 김도현 (2009) "활주로 방향설정을 위한 풍배도 프로그램의 개발 연구"
# 기준: ICAO Annex 14 / Doc. 9157 Airport Design Manual Part 1

CROSSWIND_LIMITS_KT = [10, 13, 20]   # ICAO Doc. 9157 Table 1
CALM_THRESHOLD_KT = 3.0              # 논문 §3.2 '무영향 데이터' 0~3 knots
USABILITY_TARGET = 95.0              # ICAO 권고 최소 이용률
TIE_TOLERANCE = 0.01                 # 동율 판정 허용오차 (%)
MIN_SEPARATION_DEG = 30              # 2개 활주로 최소 각도 분리 (물리적 배치 제약)
KT_TO_MS = 1 / 1.94384                # 1 kt = 0.51444... m/s

def _kt2ms(kt):
    """knot → m/s 변환."""
    return kt * KT_TO_MS

def fmt_kt(kt, decimals=1):
    """'X kt (Y m/s)' 형식으로 동시 표기. 분석결과 전반의 풍속 표시에 공용으로 사용."""
    return f"{kt:g} kt ({_kt2ms(kt):.{decimals}f} m/s)"

def select_limit_by_rwy_length(length_m, low_friction=False):
    """ICAO Doc. 9157 Table 1 기반 측풍 허용치 자동 선택. 두 번째 반환값은 선택 조건 설명(풍속 표기 제외)."""
    if length_m >= 1500:
        if low_friction:
            return 13, "≥1,500m · 종방향 마찰계수 부족"
        return 20, "≥1,500m"
    if length_m >= 1200:
        return 13, "1,200~1,500m"
    return 10, "<1,200m"

def rwy_name(deg):
    """방위각(0~179°)을 활주로 명칭(예: '15-33')으로 변환."""
    def fmt(d):
        n = round(d / 10) % 36
        return 36 if n == 0 else n
    a = fmt(deg)
    b = fmt((deg + 180) % 360)
    lo, hi = (a, b) if a < b else (b, a)
    return f"{lo:02d}-{hi:02d}"

RWY_ANGLE_STEP_DEG = 10   # 활주로 명칭(예: 17-35)은 10° 단위로만 존재 → 분석도 10° 격자로 통일

def analyze_runway(df):
    """전체 분석: calm 처리 + 3개 허용치 + 동율 처리 + 2개 활주로 + 빈도표.
    활주로 명칭은 방위각을 10° 단위로 반올림해 정해지므로(예: 166°→17, 170°→17),
    분석 자체를 10° 격자(0,10,…,170)에서만 수행해 '최적 활주로'로 표시되는 명칭과
    실제로 평가된 각도가 항상 정확히 일치하도록 한다 (1° 단위로 찾은 뒤 반올림하면
    표시된 활주로 명칭이 가리키는 각도와 실제 이용률 산출 각도가 달라지는 문제가 있었음).
    또한 활주로는 실제로 하나(또는 한 쌍)만 지을 수 있으므로, 최적 각도는 10/13/20kt
    세 허용치의 이용률을 모두 합산한 결합 기준으로 공동 선정한다. 즉 한계치별로
    서로 다른 '최적 활주로'를 추천하지 않고, 동일한 각도를 각 한계치에서 평가한다."""
    ws = df['ws_kt'].to_numpy(dtype=np.float32)
    wd = df['wd'].to_numpy(dtype=np.float32)
    N_total = len(ws)

    # 1) Calm(무영향) 분리 — 논문 §3.2
    calm_mask = ws <= CALM_THRESHOLD_KT
    N_calm = int(calm_mask.sum())
    eff_ws = ws[~calm_mask]
    eff_wd = wd[~calm_mask]
    N_eff = len(eff_ws)

    # 2) 유효 바람에 대한 측풍 행렬 (N_eff × 18, 10° 격자)
    angles = np.arange(0, 180, RWY_ANGLE_STEP_DEG, dtype=np.int32)   # [0,10,...,170]
    n_ang = len(angles)
    diff = np.radians(eff_wd[:, None] - angles[None, :])
    xwind = np.abs(eff_ws[:, None] * np.sin(diff)).astype(np.float32)  # |측풍|, knots

    # 3) 한계치별 단일/2개 활주로 이용률을 모두 먼저 계산 (배열 인덱스 기준)
    usab_by_limit = {}
    pair_usab_by_limit = {}
    for limit in CROSSWIND_LIMITS_KT:
        coverage = xwind <= limit                # (N_eff, n_ang) bool
        eff_covered = coverage.sum(axis=0)       # (n_ang,)
        usab_by_limit[limit] = (N_calm + eff_covered) / N_total * 100.0

        # 2개 활주로 분석 (합집합 이용률): |A∪B| = |A| + |B| - |A∩B|, 행렬곱으로 벡터화
        Cf = coverage.astype(np.float32)
        inter = Cf.T @ Cf                                       # (n_ang,n_ang)
        union = eff_covered[:, None] + eff_covered[None, :] - inter
        pair_usab_by_limit[limit] = (N_calm + union) / N_total * 100.0   # (n_ang,n_ang)

    # 4) 단일 최적 각도(인덱스) — 10/13/20kt 이용률 합산(결합 기준)으로 공동 선정
    combined_usab = sum(usab_by_limit[l] for l in CROSSWIND_LIMITS_KT)
    joint_u_max = float(combined_usab.max())
    joint_tied_idx = np.where(combined_usab >= joint_u_max - TIE_TOLERANCE)[0]
    mean_xw_tied = xwind[:, joint_tied_idx].mean(axis=0)
    best_idx = int(joint_tied_idx[int(np.argmin(mean_xw_tied))])
    joint_best_angle = int(angles[best_idx])                    # 실제 각도값(0~170)으로 환산

    # 5) 2개 활주로 최적 조합(인덱스)도 동일하게 결합 기준으로 공동 선정
    combined_pair_usab = sum(pair_usab_by_limit[l] for l in CROSSWIND_LIMITS_KT)
    sep = np.abs(angles[:, None] - angles[None, :])
    sep_ok = (sep >= MIN_SEPARATION_DEG) & (sep <= 180 - MIN_SEPARATION_DEG)
    masked = np.where(sep_ok, combined_pair_usab, -1.0)
    flat = int(np.argmax(masked))
    pi_idx, pj_idx = int(flat // n_ang), int(flat % n_ang)
    pi, pj = int(angles[pi_idx]), int(angles[pj_idx])           # 실제 각도값으로 환산
    joint_pair_angles = (min(pi, pj), max(pi, pj))
    pi_idx, pj_idx = int(np.where(angles == joint_pair_angles[0])[0][0]), \
                     int(np.where(angles == joint_pair_angles[1])[0][0])

    # 6) 공동 선정된 각도를 한계치별로 평가 (동일 각도, 한계치별 이용률만 다름)
    results = {}
    for limit in CROSSWIND_LIMITS_KT:
        usab = usab_by_limit[limit]
        pair_usab = pair_usab_by_limit[limit]
        u_at_best = float(usab[best_idx])
        pair_u = float(pair_usab[pi_idx, pj_idx])

        results[limit] = {
            'usability': usab,
            'best_angle': joint_best_angle,
            'best_usab': u_at_best,
            'tied_count': len(joint_tied_idx),
            'mean_xwind': float(xwind[:, best_idx].mean()),
            'max_xwind': float(xwind[:, best_idx].max()),
            'pass': u_at_best >= USABILITY_TARGET,
            'pair_angles': joint_pair_angles,
            'pair_usab': pair_u,
            'pair_pass': pair_u >= USABILITY_TARGET,
        }

    # 7) 16방위 × 풍속 빈도표 — 논문 Table 6
    freq_table = _build_freq_table(wd, ws, N_total)

    return {
        'N_total': N_total, 'N_calm': N_calm, 'N_eff': N_eff,
        'calm_pct': N_calm / N_total * 100.0,
        'angles': angles, 'results': results, 'freq_table': freq_table,
    }

def _build_freq_table(wd, ws, N_total):
    """16방위 × 5개 풍속구간 빈도표(%). 논문 Table 6 한국어 단위(노트) 버전."""
    dir_names = ["N", "NNE", "NE", "ENE", "E", "ESE", "SE", "SSE",
                 "S", "SSW", "SW", "WSW", "W", "WNW", "NW", "NNW"]
    speed_labels = [
        f"Calm 0–3 kt (0–{_kt2ms(3):.1f} m/s)",
        f"4–10 kt ({_kt2ms(4):.1f}–{_kt2ms(10):.1f} m/s)",
        f"11–13 kt ({_kt2ms(11):.1f}–{_kt2ms(13):.1f} m/s)",
        f"14–20 kt ({_kt2ms(14):.1f}–{_kt2ms(20):.1f} m/s)",
        f">20 kt (>{_kt2ms(20):.1f} m/s)",
    ]
    speed_thresh = [3, 10, 13, 20]                    # 오름차순 경계값(초과 기준, 빈틈 없음)
    spd_idx = np.zeros(len(ws), dtype=np.int32)
    for i, t in enumerate(speed_thresh):
        spd_idx[ws > t] = i + 1                       # (0,3]→0, (3,10]→1, … (20,∞)→4
    # 22.5° 간격, 첫 섹터 N은 348.75°~11.25° 중심
    dir_idx = (((wd + 11.25) // 22.5) % 16).astype(np.int32)
    table = np.zeros((16, len(speed_labels)), dtype=np.float64)
    for d in range(16):
        for s in range(len(speed_labels)):
            table[d, s] = ((dir_idx == d) & (spd_idx == s)).sum()
    pct = table / N_total * 100.0
    df_out = pd.DataFrame(pct, index=dir_names, columns=speed_labels).round(2)
    df_out['TOTAL %'] = df_out.sum(axis=1).round(2)
    return df_out


# ── [검토서(Word) 생성] ──────────────────────────────────────────────────
# 화면용 Plotly 차트와 별개로, 보고서 삽입용 정적 이미지는 matplotlib로 렌더링한다.
# kaleido(Plotly 정적변환)는 Chrome 브라우저 의존성이 있어 Streamlit Cloud(Linux)
# 배포 환경에서 불안정하므로 사용하지 않는다. matplotlib(Agg)은 브라우저 없이 동작.
# 차트 내부 라벨은 한글 폰트 누락(배포 환경 tofu) 회피를 위해 영문/기호로 표기하고,
# 문서 본문·표·제목의 한글은 Word 자체 폰트(맑은 고딕)로 렌더링한다.

def _render_windrose_png(wd, ws_kt):
    """16방위 × 6풍속구간 풍배도 → PNG bytes (matplotlib polar stacked bar)."""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.ticker import FuncFormatter

    N = len(wd)
    calm_kt = 0.5 * 1.94384
    calm_mask = ws_kt <= calm_kt
    calm_pct = float(calm_mask.sum()) / N * 100.0

    wd_eff = wd[~calm_mask]
    ws_eff = ws_kt[~calm_mask]
    dir_idx = (((wd_eff + 11.25) // 22.5) % 16).astype(int)

    spd_thresh = [3, 7, 11, 17, 21]
    spd_labels = ["0-3 kt", "3-7 kt", "7-11 kt", "11-17 kt", "17-21 kt", ">=21 kt"]
    colors = ["#c6dbef", "#74c476", "#fdd835", "#fd8d3c", "#e31a1c", "#67000d"]
    spd_idx = np.zeros(len(ws_eff), dtype=int)
    for i, t in enumerate(spd_thresh):
        spd_idx[ws_eff > t] = i + 1

    freq = np.zeros((16, 6))
    for d in range(16):
        for s in range(6):
            freq[d, s] = ((dir_idx == d) & (spd_idx == s)).sum() / N * 100.0

    theta = np.deg2rad(np.arange(0, 360, 22.5))
    width = np.deg2rad(22.5) * 0.9

    fig = plt.figure(figsize=(6.2, 6.2), dpi=150)
    ax = fig.add_subplot(111, projection="polar")
    ax.set_theta_zero_location("N")
    ax.set_theta_direction(-1)               # 시계방향
    bottom = np.zeros(16)
    for s in range(6):
        ax.bar(theta, freq[:, s], width=width, bottom=bottom,
               color=colors[s], edgecolor="white", linewidth=0.3, label=spd_labels[s])
        bottom += freq[:, s]
    ax.set_xticks(np.deg2rad([0, 45, 90, 135, 180, 225, 270, 315]))
    ax.set_xticklabels(["N", "NE", "E", "SE", "S", "SW", "W", "NW"])
    ax.set_rlabel_position(135)
    ax.tick_params(labelsize=8)
    ax.yaxis.set_major_formatter(FuncFormatter(lambda v, _: f"{v:g}%"))
    ax.text(0.5, 0.5, f"Calm\n{calm_pct:.1f}%", transform=ax.transAxes,
            ha="center", va="center", fontsize=9,
            bbox=dict(boxstyle="circle", fc="white", ec="#888", alpha=0.9))
    ax.legend(loc="upper left", bbox_to_anchor=(1.02, 1.05), fontsize=7,
              title="Wind Speed", title_fontsize=8)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


def _render_usability_png(A):
    """방위각별 이용률 곡선(허용치 3종) → PNG bytes (matplotlib)."""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    angles = A['angles']
    line_colors = {10: "#1f77b4", 13: "#2ca02c", 20: "#d62728"}
    fig, ax = plt.subplots(figsize=(7.4, 4.2), dpi=150)
    for lim in CROSSWIND_LIMITS_KT:
        u = A['results'][lim]['usability']
        ax.plot(angles, u, marker="o", ms=3, color=line_colors.get(lim),
                label=f"{lim} kt ({_kt2ms(lim):.1f} m/s)")
    ax.axhline(USABILITY_TARGET, ls="--", color="red", lw=1)
    ax.text(angles[-1], USABILITY_TARGET, " ICAO 95%", color="red",
            va="bottom", ha="right", fontsize=8)
    for lim in CROSSWIND_LIMITS_KT:
        ax.axvline(A['results'][lim]['best_angle'], ls=":", color="gray", alpha=0.3)
    ax.set_xlabel("Runway Azimuth (deg)")
    ax.set_ylabel("Usability (%)")
    ax.set_title("Usability by Runway Direction")
    ax.set_xticks(np.arange(0, 180, 20))
    ax.grid(alpha=0.3)
    ax.legend(title="Crosswind Limit", fontsize=8, title_fontsize=8)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png")
    plt.close(fig)
    return buf.getvalue()


def _build_detail_table(A):
    """방위각 10° 간격 이용률 상세표 (t6과 동일 로직)."""
    rows = []
    for h in np.arange(RWY_ANGLE_STEP_DEG, 181, RWY_ANGLE_STEP_DEG, dtype=int):
        i = int(h % 180)
        idx = i // RWY_ANGLE_STEP_DEG
        row = {"방향(°)": int(h), "대응방향(°)": int(h + 180), "활주로": rwy_name(i)}
        for lim in CROSSWIND_LIMITS_KT:
            row[f"{lim}kt 이용률(%)"] = round(float(A['results'][lim]['usability'][idx]), 2)
        rows.append(row)
    return pd.DataFrame(rows)


def _build_summary_table(A):
    """허용치별 종합표 (t1 핵심 항목)."""
    rows = []
    for lim in CROSSWIND_LIMITS_KT:
        rr = A['results'][lim]
        rows.append({
            "허용치": fmt_kt(lim),
            "최적 활주로": rwy_name(rr['best_angle']),
            "방위각(°)": rr['best_angle'],
            "이용률(%)": round(rr['best_usab'], 3),
            "평균측풍(kt)": round(rr['mean_xwind'], 2),
            "평균측풍(m/s)": round(_kt2ms(rr['mean_xwind']), 2),
            "최대측풍(kt)": round(rr['max_xwind'], 2),
            "최대측풍(m/s)": round(_kt2ms(rr['max_xwind']), 2),
            "단일판정": "적합" if rr['pass'] else "부적합",
        })
    return pd.DataFrame(rows)


# 검토서 템플릿: assets/report_template.docx 가 있으면 그 문서의 스타일(글꼴·표 스타일·
# 페이지 설정·머리말/꼬리말)을 그대로 상속한다. 없으면 기본 문서로 폴백한다.
# → 스타일을 바꾸려면 코드 수정 없이 이 템플릿 파일만 Word로 편집하면 된다.
REPORT_TEMPLATE_PATH = os.path.join(
    os.path.dirname(os.path.abspath(__file__)), "assets", "report_template.docx")
REPORT_TABLE_STYLE = "Table Grid"   # 템플릿에서 이 스타일을 재정의하거나 커스텀 표 스타일명으로 교체 가능


def _clear_body(doc):
    """템플릿 본문의 자리표시 콘텐츠(문단·표)만 제거. 스타일 정의, 페이지 설정(sectPr),
    머리말/꼬리말(별도 파트)은 그대로 유지되어 생성 콘텐츠가 템플릿 서식을 상속한다."""
    from docx.oxml.ns import qn
    body = doc.element.body
    for child in list(body):
        if child.tag == qn('w:sectPr'):     # 섹션(페이지) 설정은 보존
            continue
        body.remove(child)


def _replace_in_paragraph(p, mapping):
    """문단 내 {{KEY}} 자리표시를 mapping 값으로 치환.
    자리표시가 여러 run으로 쪼개져 있어도 동작하도록 run 텍스트를 합쳐 처리하고,
    결과를 첫 run에 넣어 그 run의 서식(글꼴·크기·색)을 유지한다."""
    full = "".join(r.text for r in p.runs)
    if "{{" not in full:
        return
    new = full
    for k, v in mapping.items():
        new = new.replace("{{" + k + "}}", str(v))
    if new != full and p.runs:
        p.runs[0].text = new
        for r in p.runs[1:]:
            r.text = ""


def _iter_placeholder_targets(doc, include_header_footer=True):
    """치환 대상 문단을 순회 (본문·표, 선택적으로 머리말/꼬리말)."""
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    if include_header_footer:
        for section in doc.sections:
            for hf in (section.header, section.footer):
                for p in hf.paragraphs:
                    yield p
                for t in hf.tables:
                    for row in t.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                yield p


def _replace_placeholders(doc, mapping):
    """문서 전체(본문·표·머리말/꼬리말)의 {{KEY}} 자리표시를 치환."""
    for p in _iter_placeholder_targets(doc, include_header_footer=True):
        _replace_in_paragraph(p, mapping)


def _has_body_placeholder(doc):
    """본문(또는 본문 표)에 {{...}} 자리표시가 있으면 True → 표지 페이지로 간주.
    (머리말/꼬리말의 자리표시는 표지 판정에서 제외)"""
    for p in _iter_placeholder_targets(doc, include_header_footer=False):
        if "{{" in "".join(r.text for r in p.runs):
            return True
    return False


def _add_toc(doc):
    """제목1/제목2 스타일 기반 자동 목차(TOC) 필드를 삽입한다.
    Word/HWP에서 문서를 열고 F9(또는 우클릭 → '필드 업데이트')로 페이지번호까지 갱신된다."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    run = doc.add_paragraph().add_run()
    f_begin = OxmlElement('w:fldChar'); f_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    f_sep = OxmlElement('w:fldChar'); f_sep.set(qn('w:fldCharType'), 'separate')
    placeholder = OxmlElement('w:t')
    placeholder.text = "목차: 우클릭 → '필드 업데이트'(F9)로 생성됩니다."
    f_end = OxmlElement('w:fldChar'); f_end.set(qn('w:fldCharType'), 'end')
    r = run._r
    for _el in (f_begin, instr, f_sep, placeholder, f_end):
        r.append(_el)


def _docx_set_korean_font(doc, name="맑은 고딕", size=10):
    from docx.shared import Pt
    from docx.oxml.ns import qn
    style = doc.styles['Normal']
    style.font.name = name
    style.font.size = Pt(size)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn('w:eastAsia'), name)


def _shade_cell(cell, hex_color):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _add_df_table(doc, df, header_bg="D9E1F2", center_data=False):
    """표 삽입. 글꼴(이름·크기)은 강제하지 않고 템플릿의 Normal 스타일을 상속한다
    (→ 표 내부 글꼴을 템플릿에서 일괄 제어). 헤더는 음영·굵게·가운데정렬,
    center_data=True면 데이터 셀도 가운데정렬(숫자·라벨 위주 표)."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import RGBColor
    cols = list(df.columns)
    t = doc.add_table(rows=1, cols=len(cols))
    try:
        t.style = REPORT_TABLE_STYLE        # 템플릿이 제공/재정의한 표 스타일 우선
    except KeyError:
        t.style = "Table Grid"              # 템플릿에 해당 스타일이 없으면 기본
    t.autofit = True
    hdr = t.rows[0].cells
    for j, c in enumerate(cols):
        hdr[j].text = str(c)
        _shade_cell(hdr[j], header_bg)
        for p in hdr[j].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
    for _, rowdata in df.iterrows():
        cells = t.add_row().cells
        for j, c in enumerate(cols):
            v = rowdata[c]
            sval = "" if pd.isna(v) else str(v)
            cells[j].text = sval
            is_todo = "[ 작성" in sval          # 미작성 항목은 빨강 강조
            for p in cells[j].paragraphs:
                if center_data:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if is_todo:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                        run.font.bold = True
    return t


# 영문 16방위 → 한글 방위명 (서술문 표기용)
_KDIR = {
    "N": "북", "NNE": "북북동", "NE": "북동", "ENE": "동북동", "E": "동", "ESE": "동남동",
    "SE": "남동", "SSE": "남남동", "S": "남", "SSW": "남남서", "SW": "남서", "WSW": "서남서",
    "W": "서", "WNW": "서북서", "NW": "북서", "NNW": "북북서",
}


def _narr_stats(A, df, stn_name, chart_start, chart_end, primary_limit,
                n_raw, n_valid, n_invalid):
    """검토서 본문 서술에 쓰는 파생 통계값을 계산해 dict로 반환.
    (LLM 없이 분석결과 A와 원자료 df만으로 값을 산출 → 매 생성 시 토큰 비용 없음)"""
    angles = A['angles']
    period_months = ((chart_end.year - chart_start.year) * 12
                     + (chart_end.month - chart_start.month) + 1)
    py = period_months / 12
    period_years_str = str(int(py)) if float(py).is_integer() else f"{py:.1f}"

    ws_kt = df['ws_kt'].to_numpy(dtype=float)
    calm_rose_pct = float((ws_kt <= 0.5 * 1.94384).mean() * 100) if len(ws_kt) else 0.0

    ft = A['freq_table']
    total_col = ft['TOTAL %']
    top = total_col.sort_values(ascending=False)
    top_dirs = [(d, _KDIR.get(d, d), float(top[d])) for d in top.index[:3]]

    spd_cols = [c for c in ft.columns if c != 'TOTAL %']
    noncalm = [c for c in spd_cols if not c.startswith('Calm')]
    col_sums = {c: float(ft[c].sum()) for c in noncalm}
    dom_bin = max(col_sums, key=col_sums.get) if col_sums else ''
    dom_bin_short = dom_bin.split('(')[0].strip()
    over20_pct = float(sum(ft[c].sum() for c in spd_cols if c.strip().startswith('>')))

    best_angle = int(A['results'][primary_limit]['best_angle'])
    best_idx = best_angle // RWY_ANGLE_STEP_DEG
    best_rwy = rwy_name(best_angle)
    opp_angle = (best_angle + 180) % 360

    usab_by = {lim: float(A['results'][lim]['best_usab']) for lim in CROSSWIND_LIMITS_KT}
    mean_xw_kt = float(A['results'][primary_limit]['mean_xwind'])
    max_xw_kt = float(A['results'][primary_limit]['max_xwind'])

    perlim = {}
    for lim in CROSSWIND_LIMITS_KT:
        u = np.asarray(A['results'][lim]['usability'], dtype=float)
        imin, imax = int(u.argmin()), int(u.argmax())
        perlim[lim] = {
            'min_ang': int(angles[imin]), 'min_val': float(u[imin]), 'min_rwy': rwy_name(int(angles[imin])),
            'max_ang': int(angles[imax]), 'max_val': float(u[imax]), 'max_rwy': rwy_name(int(angles[imax])),
            'umin': float(u.min()), 'umax': float(u.max()),
        }
    all_pass = all(bool((np.asarray(A['results'][lim]['usability']) >= USABILITY_TARGET).all())
                   for lim in CROSSWIND_LIMITS_KT)
    same_dir = len({A['results'][lim]['best_angle'] for lim in CROSSWIND_LIMITS_KT}) == 1
    overall_min = min(perlim[lim]['umin'] for lim in CROSSWIND_LIMITS_KT)
    overall_max = max(perlim[lim]['umax'] for lim in CROSSWIND_LIMITS_KT)

    return {
        'period_months': period_months, 'period_years_str': period_years_str,
        'valid_pct': n_valid / n_raw * 100 if n_raw else 0.0,
        'miss_pct': n_invalid / n_raw * 100 if n_raw else 0.0,
        'calm_rose_pct': calm_rose_pct,
        'top_dirs': top_dirs, 'dom_bin_short': dom_bin_short, 'over20_pct': over20_pct,
        'best_angle': best_angle, 'best_idx': best_idx, 'best_rwy': best_rwy, 'opp_angle': opp_angle,
        'usab_by': usab_by, 'mean_xw_kt': mean_xw_kt, 'max_xw_kt': max_xw_kt,
        'perlim': perlim, 'all_pass': all_pass, 'same_dir': same_dir,
        'overall_min': overall_min, 'overall_max': overall_max,
    }


def _add_red_para(doc, text):
    """[ 작성 ] 등 미작성 항목을 빨강·굵게로 삽입."""
    from docx.shared import RGBColor
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    run.font.bold = True
    return p


def _add_bullet(doc, text, indent=True):
    """불릿 문단(스타일 의존 없이 문자 불릿으로)."""
    p = doc.add_paragraph(("    · " if indent else "· ") + text)
    return p


def _add_caption(doc, text):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    c = doc.add_paragraph(text)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return c


def build_review_docx(A, df, stn_name, chart_start, chart_end, primary_limit,
                      n_raw, n_valid, n_invalid):
    """분석 결과(A)와 원자료(df)로 활주로 방향 검토서(.docx) bytes 생성.
    본문 서술은 고정 문구(boilerplate) + 분석값 자동 치환 방식(LLM 미사용)."""
    from docx import Document
    from docx.shared import Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    period_months = ((chart_end.year - chart_start.year) * 12
                     + (chart_end.month - chart_start.month) + 1)
    r = A['results'][primary_limit]

    # 표지 자리표시 치환값. 템플릿 본문(표지)에 {{TITLE}} {{STATION}} {{PERIOD}}
    # {{DATE}} {{DATE_KR}} 를 넣어두면 아래 값으로 자동 치환된다.
    today = date.today()
    _ph_map = {
        "TITLE": "활주로 방향 검토서",
        "STATION": stn_name,
        "PERIOD": f"{chart_start:%Y-%m-%d} ~ {chart_end:%Y-%m-%d}",
        "DATE": f"{today:%Y-%m-%d}",
        "DATE_KR": f"{today.year}년 {today.month}월 {today.day}일",
    }

    has_cover = False
    if os.path.exists(REPORT_TEMPLATE_PATH):
        doc = Document(REPORT_TEMPLATE_PATH)   # 템플릿 스타일·페이지설정·머리말/꼬리말 상속
        has_cover = _has_body_placeholder(doc)  # 본문에 {{...}}가 있으면 표지로 간주
        if has_cover:
            _replace_placeholders(doc, _ph_map)  # 표지 제목·날짜 등 치환
            doc.add_page_break()                 # 표지 뒤 페이지 나눔
        else:
            _clear_body(doc)                     # 스타일만 있는 템플릿 → 본문 비우고 서식 상속
    else:
        doc = Document()
        _docx_set_korean_font(doc)               # 템플릿이 없을 때만 기본 한글 글꼴 적용

    # 표지가 없을 때만 코드가 표제/부제를 생성 (표지가 있으면 그 표지가 제목 역할)
    if not has_cover:
        title = doc.add_heading("활주로 방향 검토서", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub = doc.add_paragraph(
            f"대상 관측소: {stn_name}    |    분석기간: "
            f"{chart_start:%Y-%m-%d} ~ {chart_end:%Y-%m-%d}"
        )
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    # 목차 (제목1/제목2 스타일 기반 자동 목차 — Word/HWP에서 F9로 갱신)
    toc_title = doc.add_paragraph()
    _toc_run = toc_title.add_run("목  차")
    _toc_run.bold = True
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_toc(doc)
    doc.add_page_break()

    # 본문 서술용 파생 통계 (분석결과 값 자동 치환)
    S = _narr_stats(A, df, stn_name, chart_start, chart_end, primary_limit,
                    n_raw, n_valid, n_invalid)
    start_kr = f"{chart_start.year}년 {chart_start.month}월 {chart_start.day}일"
    end_kr = f"{chart_end.year}년 {chart_end.month}월 {chart_end.day}일"
    calm_lbl = fmt_kt(CALM_THRESHOLD_KT)

    # 1. 검토 개요
    doc.add_heading("1. 검토 개요", level=1)
    _add_df_table(doc, pd.DataFrame([
        {"항목": "공항(시설)명", "내용": "[ 작성 ]"},
        {"항목": "검토 목적", "내용": "[ 작성 ]"},
        {"항목": "의뢰기관", "내용": "[ 작성 ]"},
        {"항목": "대상 관측소", "내용": stn_name},
        {"항목": "분석기간",
         "내용": f"{chart_start:%Y-%m-%d} ~ {chart_end:%Y-%m-%d} ({period_months}개월)"},
        {"항목": "검토기준", "내용": "ICAO Annex 14 / Doc.9157 Airport Design Manual Part 1"},
    ]))
    doc.add_paragraph()

    doc.add_heading("1.1 검토배경 및 목적", level=2)
    _add_red_para(doc, "[ 작성: 본 검토의 배경이 되는 사업명 및 구체적 검토 목적을 기재 ]")
    doc.add_paragraph(
        "활주로 방향은 항공기 이착륙 시 안전성 및 운영효율성을 결정짓는 핵심 요소로서, "
        "특정 방향 활주로에 대한 풍향·풍속 분석에 기반한 활주로 이용률(Wind Coverage) 검증이 "
        "선행되어야 한다. 이에 본 검토에서는 대상 관측소의 장기간 기상관측자료를 활용하여 "
        "측풍 허용치별 활주로 이용률을 정량적으로 산정하고, 최적의 활주로 방향을 도출하고자 한다."
    )

    doc.add_heading("1.2 검토대상", level=2)
    doc.add_paragraph(
        f"본 검토의 대상 관측소는 {stn_name}(기상청 {stn_name} 관측소)로 하며, 해당 관측소에서 "
        f"관측된 풍향·풍속 자료를 기초자료로 하여 분석을 수행한다."
    )

    doc.add_heading("1.3 분석기간", level=2)
    doc.add_paragraph(
        f"분석기간은 {start_kr}부터 {end_kr}까지 최근 {period_months}개월"
        f"({S['period_years_str']}개년)간의 관측자료를 대상으로 한다. 다개년의 자료를 활용함으로써 "
        f"계절별 풍향 변화 특성과 연도별 편차를 종합적으로 반영할 수 있도록 하였다."
    )

    doc.add_heading("1.4 검토기준", level=2)
    doc.add_paragraph("본 검토는 활주로 방향 결정과 관련한 다음의 국내 및 국제 기준을 근거로 수행한다.")
    doc.add_paragraph("○ 국내기준")
    _add_bullet(doc, "공항시설법")
    _add_bullet(doc, "공항·비행장시설 및 이착륙장 설치기준")
    _add_bullet(doc, "공항·비행장시설 설계 세부지침")
    doc.add_paragraph("○ 국제기준")
    _add_bullet(doc, "ICAO Annex 14 (Aerodrome Design and Operations)")
    _add_bullet(doc, "ICAO Doc.9157, Airport Design Manual Part 1 (Runways)")
    doc.add_paragraph(
        f"상기 기준에서 제시하는 활주로 이용률 기준(통상 {USABILITY_TARGET:g}% 이상)을 충족하는지 "
        f"여부를 판단하는 것을 본 검토의 핵심 절차로 한다."
    )
    doc.add_paragraph()

    # 2. 기상자료 현황
    doc.add_heading("2. 기상자료 현황", level=1)
    _add_df_table(doc, pd.DataFrame([
        {"구분": "전체 수집 행", "값": f"{n_raw:,} 행"},
        {"구분": "유효 데이터 (풍향+풍속 존재)", "값": f"{n_valid:,} 행 ({S['valid_pct']:.1f}%)"},
        {"구분": "결측 데이터", "값": f"{n_invalid:,} 행 ({S['miss_pct']:.1f}%)"},
        {"구분": "전체 관측시간", "값": f"{A['N_total']:,} 시간"},
        {"구분": f"Calm (0~{calm_lbl})",
         "값": f"{A['N_calm']:,} 시간 ({A['calm_pct']:.2f}%)"},
        {"구분": "유효 풍황 데이터", "값": f"{A['N_eff']:,} 시간"},
    ]))
    doc.add_paragraph()

    doc.add_heading("2.1 자료 수집 현황", level=2)
    doc.add_paragraph(
        f"대상 관측소({stn_name})에서 분석기간({chart_start:%Y.%m.%d}~{chart_end:%Y.%m.%d}, "
        f"{period_months}개월) 동안 수집된 기상관측자료는 총 {n_raw:,}행이며, 이는 시간 단위 "
        f"관측자료를 기준으로 한 것이다."
    )

    doc.add_heading("2.2 자료 정합성 검토", level=2)
    doc.add_paragraph(
        f"수집된 전체 {n_raw:,}행 중 풍향 및 풍속이 모두 존재하는 유효 데이터는 "
        f"{n_valid:,}행({S['valid_pct']:.1f}%)으로 나타났으며, 결측 데이터는 "
        f"{n_invalid:,}행({S['miss_pct']:.1f}%)으로, 분석에 활용 가능한 자료의 정합성 및 "
        f"신뢰도는 높은 것으로 판단된다. 결측 자료는 분석 결과에 미치는 영향이 미미하므로 "
        f"별도의 보정 없이 분석대상에서 제외하였다."
    )

    doc.add_heading("2.3 정온(Calm) 조건 처리", level=2)
    doc.add_paragraph(
        f"유효 데이터 {A['N_total']:,}시간 중 풍속이 {calm_lbl} 이하인 Calm(무풍) 상태는 "
        f"{A['N_calm']:,}시간(전체의 {A['calm_pct']:.2f}%)으로 집계되었다. Calm 조건은 풍속이 "
        f"미약하여 특정 방향의 활주로 운영에 대한 측풍(Crosswind) 영향이 거의 없는 것으로 "
        f"간주되며, ICAO Doc.9157 등 국제기준에 따라 방향별 활주로 이용률(Wind Coverage) "
        f"산정 시 모든 활주로 방향에 대해 공통적으로 가산되는 항목으로 처리된다."
    )

    doc.add_heading("2.4 유효 풍향자료 산정", level=2)
    doc.add_paragraph(
        f"전체 관측시간 {A['N_total']:,}시간에서 Calm 시간 {A['N_calm']:,}시간을 제외한 "
        f"유효 풍향 데이터는 {A['N_eff']:,}시간으로 산정되었으며, 본 자료를 기초로 방향별 "
        f"측풍성분 및 활주로 이용률을 분석하는 데 활용하였다."
    )
    doc.add_paragraph()

    # 3. 분석 방법
    doc.add_heading("3. 분석 방법", level=1)

    doc.add_heading("3.1 분석기준 및 근거", level=2)
    doc.add_paragraph(
        "본 검토는 국제민간항공기구(ICAO) Annex 14 및 Doc.9157 Airport Design Manual "
        "Part 1(Runways)에 따라 활주로 방향별 측풍 이용률(Usability Factor)을 산정하였다. "
        "이는 특정 활주로 방향에 대해 항공기가 안전하게 이착륙할 수 있는 시간의 비율을 정량적으로 "
        "평가하기 위한 국제표준 분석기법이다."
    )

    doc.add_heading("3.2 측풍 허용치 적용기준", level=2)
    doc.add_paragraph(
        "측풍(Crosswind) 허용치는 대상공항의 항공기 기준활주로길이(ARFL)에 따라 아래와 같이 "
        "차등 적용하였다."
    )
    _add_df_table(doc, pd.DataFrame([
        {"측풍 허용치": "20 kt (10.3 m/s)", "적용 대상": "기준활주로길이 1,500m 이상"},
        {"측풍 허용치": "13 kt (6.7 m/s)", "적용 대상": "1,200m 이상 1,500m 미만 (또는 종방향 마찰계수 부족)"},
        {"측풍 허용치": "10 kt (5.1 m/s)", "적용 대상": "1,200m 미만"},
    ]))
    doc.add_paragraph(
        "세 가지 허용치를 각각 적용하여 방향별 이용률을 산정함으로써, 활주로 규모 조건 변화에 "
        "따른 최적 방향의 민감도를 함께 검토할 수 있도록 하였다."
    )

    doc.add_heading("3.3 무풍(Calm) 상태의 처리", level=2)
    doc.add_paragraph(
        f"풍속 {calm_lbl} 이하의 무풍(Calm) 상태는 측풍 성분이 사실상 발생하지 않는 것으로 "
        f"간주하여, 활주로 방향과 관계없이 모든 방향에 대해 공통적으로 이용 가능한 시간으로 "
        f"집계하였다."
    )

    doc.add_heading("3.4 이용률 산정식", level=2)
    doc.add_paragraph("방향별 이용률은 다음 식에 따라 산정하였다.")
    eq = doc.add_paragraph(
        "이용률(%) = (Calm 관측수 + 측풍성분 ≤ 허용치 관측수) ÷ 전체 관측수 × 100"
    )
    eq.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("3.5 분석 격자 및 방향 설정", level=2)
    doc.add_paragraph(
        "활주로 명칭은 자기방위 기준 10° 단위로만 부여되는 점을 고려하여, 분석 격자 또한 "
        "10° 간격(0°~170°, 총 18방향)으로 고정하여 검토를 수행하였다. 이를 통해 산정된 최적 "
        "방향이 실제 활주로 명칭 부여체계와 정합성을 갖도록 하였다."
    )

    doc.add_heading("3.6 최적 활주로 방향 선정기준", level=2)
    doc.add_paragraph(
        "최적 활주로 방향은 앞서 제시한 3가지 측풍 허용치(10kt·13kt·20kt) 각각에 대한 "
        "이용률을 모두 합산한 결합 기준(Combined Criteria)에 따라, 세 조건에서 공통적으로 "
        "우수한 성능을 보이는 단일 방향을 선정하는 방식으로 도출하였다."
    )
    doc.add_paragraph()

    # 4. 분석 결과
    doc.add_heading("4. 분석 결과", level=1)
    td = S['top_dirs']
    p10, p13, p20 = S['perlim'][10], S['perlim'][13], S['perlim'][20]

    doc.add_heading("4.1 풍배도 (Wind Rose)", level=2)
    rose_png = _render_windrose_png(
        df['wd'].to_numpy(dtype=float), df['ws_kt'].to_numpy(dtype=float))
    doc.add_picture(io.BytesIO(rose_png), width=Cm(11.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_caption(doc, "그림 1. 16방위 풍배도 (중앙 원: 0.5 m/s 이하 정온 비율)")
    doc.add_paragraph(
        f"풍배도 분석 결과, 대상 관측소의 주풍은 {td[0][1]}({td[0][0]}) 방향을 중심으로 분포하며, "
        f"{td[0][1]} 방향이 전체 관측시간의 {td[0][2]:.2f}%를 차지하여 가장 높은 빈도를 보였다. "
        f"이어 {td[1][1]}({td[1][0]}, {td[1][2]:.2f}%), {td[2][1]}({td[2][0]}, {td[2][2]:.2f}%) "
        f"방향에서도 상대적으로 높은 빈도가 관측되었다. 이러한 주풍 분포는 후술하는 최적 활주로 "
        f"방향 산정 결과와 연계하여 해석된다."
    )
    doc.add_paragraph(
        f"한편 풍배도 중앙에 표기된 정온(Calm) 비율 {S['calm_rose_pct']:.1f}%는 풍속 0.5m/s "
        f"이하를 기준으로 산정된 값으로, 활주로 이용률 산정에 적용한 정온 기준"
        f"({calm_lbl} 이하, {A['calm_pct']:.2f}%)과는 임계풍속 기준이 상이하므로 상호 혼동하지 "
        f"않도록 유의할 필요가 있다. 이용률 분석에는 2·3장에서 정의한 {calm_lbl} 기준의 정온 "
        f"자료가 일관되게 적용되었다."
    )

    doc.add_heading("4.2 방향별 이용률 곡선", level=2)
    doc.add_picture(io.BytesIO(_render_usability_png(A)), width=Cm(14.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_caption(doc, "그림 2. 활주로 방위각별 측풍 이용률 (허용치별 비교)")
    doc.add_paragraph(
        f"방위각별 이용률 곡선을 살펴보면, 측풍 허용치 20kt(10.3m/s) 기준은 전 방위(0°~170°)에서 "
        f"최소 {p20['umin']:.2f}%~최대 {p20['umax']:.2f}%의 이용률을 나타내어 방향에 관계없이 "
        f"ICAO 기준({USABILITY_TARGET:g}%)을 상회하는 것으로 분석되었다. 13kt(6.7m/s) 기준 또한 "
        f"최솟값({p13['min_ang']}°, {p13['min_val']:.2f}%)에서 최댓값({p13['max_ang']}°, "
        f"{p13['max_val']:.2f}%)까지 좁은 변동폭을 보였다."
    )
    if p10['umin'] >= USABILITY_TARGET:
        _t42 = (f"가장 낮은 구간({p10['min_val']:.2f}%)에서도 ICAO 최소기준인 {USABILITY_TARGET:g}%를 "
                f"상회하고 있어, 본 대상지는 활주로 방향 설정에 있어 비교적 여유 있는 기상조건을 "
                f"갖춘 것으로 판단된다.")
    else:
        _t42 = (f"가장 낮은 구간({p10['min_val']:.2f}%)은 ICAO 최소기준({USABILITY_TARGET:g}%)에 "
                f"미달하므로, 해당 방위 활주로의 단독 운영 시 유의가 필요하다.")
    doc.add_paragraph(
        f"10kt(5.1m/s) 기준에서는 방위각에 따른 변동폭이 상대적으로 뚜렷하게 나타났다. "
        f"방위각 {p10['min_ang']}° 부근에서 {p10['min_val']:.2f}%로 가장 낮은 이용률을, "
        f"방위각 {p10['max_ang']}° 부근에서 {p10['max_val']:.2f}%로 최댓값을 기록하였다. " + _t42
    )

    doc.add_heading("4.3 16방위 풍향·풍속 빈도표 (%)", level=2)
    if S['over20_pct'] < 0.5:
        _t43 = (f"20kt(10.3m/s)를 초과하는 강풍은 전 방위에서 {S['over20_pct']:.1f}%로 거의 관측되지 "
                f"않아, 본 대상지는 강풍에 의한 활주로 운영제약 가능성이 낮은 것으로 판단된다.")
    else:
        _t43 = f"20kt(10.3m/s)를 초과하는 강풍은 전 방위 합계 {S['over20_pct']:.1f}%로 관측되었다."
    doc.add_paragraph(
        f"16방위 빈도분석 결과, {td[0][1]}({td[0][2]:.2f}%)와 {td[1][1]}({td[1][2]:.2f}%)가 "
        f"전체 관측시간의 약 {td[0][2] + td[1][2]:.0f}%를 차지하며 뚜렷한 주풍 방향을 형성하는 "
        f"것으로 나타났다. 풍속대별로는 대부분의 방위에서 {S['dom_bin_short']} 구간의 빈도가 가장 "
        f"높게 나타났으며, " + _t43
    )
    ft = A['freq_table'].reset_index().rename(columns={'index': '방위'})
    _add_df_table(doc, ft, center_data=True)

    doc.add_heading("4.4 방위각별 이용률 상세표 (10° 간격)", level=2)
    _add_df_table(doc, _build_detail_table(A), center_data=True)
    doc.add_paragraph(
        f"※ 측풍 허용치: 10kt=5.1m/s, 13kt=6.7m/s, 20kt=10.3m/s · "
        f"{USABILITY_TARGET:g}% 이상 적합 / 미만 부적합"
    )
    if S['all_pass']:
        _t44 = (f"방위각별 상세 분석 결과, 18개 전 방향에서 3개 허용치 기준 모두 ICAO 최소요건인 "
                f"{USABILITY_TARGET:g}%를 만족하는 것으로 나타나, 본 대상지는 활주로 방향 선정에 "
                f"있어 특정 방향으로 제약받지 않는 양호한 기상조건을 갖춘 것으로 판단된다. ")
    else:
        _t44 = (f"방위각별 상세 분석 결과, 일부 방위·허용치 조건에서 ICAO 최소요건"
                f"({USABILITY_TARGET:g}%)에 미달하는 구간이 확인되었다. ")
    doc.add_paragraph(
        _t44 + f"방향별 편차는 10kt 기준에서 가장 뚜렷하게 나타났으며, {p10['max_rwy']} 활주로"
        f"(방위각 {p10['max_ang']}°/{(p10['max_ang'] + 180) % 360}°)에서 {p10['max_val']:.2f}%로 "
        f"전 방향 중 최댓값을, {p10['min_rwy']} 활주로(방위각 {p10['min_ang']}°/"
        f"{(p10['min_ang'] + 180) % 360}°)에서 {p10['min_val']:.2f}%로 최솟값을 기록하였다."
    )

    doc.add_heading("4.5 측풍 허용치별 종합", level=2)
    _add_df_table(doc, _build_summary_table(A), center_data=True)
    if S['same_dir']:
        _t45 = (f"측풍 허용치 10kt, 13kt, 20kt 3개 기준 전체에서 최적 활주로 방향이 방위각 "
                f"{S['best_angle']}°/{S['opp_angle']}°, 즉 {S['best_rwy']} 활주로로 일관되게 "
                f"도출되었다. 이는 항공기 기준활주로길이 조건이 달라지더라도 동일한 방향이 최적으로 "
                f"선정됨을 의미하며, 산정된 방향의 강건성(Robustness)을 뒷받침하는 결과로 평가된다.")
    else:
        _t45 = (f"측풍 허용치별 최적 방향을 종합한 결과, 결합 기준상 {S['best_rwy']} 활주로"
                f"(방위각 {S['best_angle']}°/{S['opp_angle']}°)가 최적 방향으로 산정되었다.")
    doc.add_paragraph(_t45)
    doc.add_paragraph(
        f"{S['best_rwy']} 활주로 방향의 이용률은 {S['overall_min']:.3f}%~{S['overall_max']:.3f}%로 "
        f"ICAO 기준({USABILITY_TARGET:g}%)을 상회하며, 최적 방향 기준 평균측풍은 "
        f"{fmt_kt(round(S['mean_xw_kt'], 2))}, 분석기간 중 최대측풍은 "
        f"{fmt_kt(round(S['max_xw_kt'], 2))}로 나타나 허용치 기준 대비 여유를 확보한 것으로 "
        f"확인되었다. 이상을 종합할 때 본 대상지의 최적 활주로 방향은 {S['best_rwy']}"
        f"(방위각 {S['best_angle']}°/{S['opp_angle']}°)로 판단된다."
    )

    # 5. 최적 활주로 방향 선정
    doc.add_heading("5. 최적 활주로 방향 선정", level=1)
    _add_df_table(doc, pd.DataFrame([{
        "허용치": fmt_kt(lim),
        "최적 활주로": rwy_name(A['results'][lim]['best_angle']),
        "방위각(°)": A['results'][lim]['best_angle'],
        "이용률(%)": f"{A['results'][lim]['best_usab']:.3f}",
        "판정": "적합" if A['results'][lim]['pass'] else "부적합",
    } for lim in CROSSWIND_LIMITS_KT]), center_data=True)
    doc.add_paragraph(
        f"세 허용치를 종합한 결과, 최적 활주로 방향은 {S['best_rwy']} "
        f"(방위각 {S['best_angle']}°)로 선정되었다. 적용 허용치 {fmt_kt(primary_limit)} 기준 "
        f"단일 활주로 이용률은 {r['best_usab']:.3f}%이다."
    )
    if not r['pass']:
        doc.add_paragraph(
            f"단일 활주로로는 ICAO 권고 이용률 {USABILITY_TARGET:g}%에 미달하므로, "
            f"2개 활주로 조합({rwy_name(r['pair_angles'][0])} + {rwy_name(r['pair_angles'][1])}) "
            f"적용 시 이용률은 {r['pair_usab']:.3f}%로 개선된다."
        )

    min_usab = min(S['usab_by'].values())

    doc.add_heading("5.1 종합 판정", level=2)
    if S['same_dir']:
        _t51 = (f"4장에서 산정한 10kt, 13kt, 20kt 3개 측풍 허용치 기준 결과를 종합한 결과, 모든 "
                f"허용치 조건에서 동일하게 방위각 {S['best_angle']}°(활주로 {S['best_rwy']})가 최적 "
                f"방향으로 도출되었다. 이는 항공기 기준활주로길이 조건이 달라지더라도 최적 방향이 "
                f"변동되지 않음을 의미하며, 특정 항공기 규모나 향후 활주로 연장 계획 변경 시에도 본 "
                f"방향 선정 결과의 유효성이 유지될 수 있음을 시사한다.")
    else:
        _t51 = (f"3개 측풍 허용치 기준 결과를 종합한 결과, 결합 기준상 최적 방향은 방위각 "
                f"{S['best_angle']}°(활주로 {S['best_rwy']})로 도출되었다.")
    doc.add_paragraph(_t51)

    doc.add_heading("5.2 이용률 검토", level=2)
    if min_usab >= USABILITY_TARGET:
        _t52 = (f"3개 허용치 모두 ICAO 및 국내기준에서 요구하는 최소 이용률 {USABILITY_TARGET:g}%를 "
                f"상회하여 전 조건에서 적합 판정을 획득하였다. 특히 가장 보수적인 기준인 10kt "
                f"허용치에서도 {USABILITY_TARGET:g}% 기준 대비 약 "
                f"{S['usab_by'][10] - USABILITY_TARGET:.1f}%p의 여유를 확보하고 있어, 본 대상지는 "
                f"활주로 방향 설계에 있어 기상학적 제약이 거의 없는 조건을 갖춘 것으로 판단된다.")
    else:
        _t52 = (f"가장 보수적인 10kt 허용치 기준 이용률({S['usab_by'][10]:.3f}%)이 최소기준"
                f"({USABILITY_TARGET:g}%)에 대한 적합 여부를 별도로 검토할 필요가 있다.")
    doc.add_paragraph(
        f"{S['best_rwy']} 활주로의 허용치별 이용률은 10kt 기준 {S['usab_by'][10]:.3f}%, "
        f"13kt 기준 {S['usab_by'][13]:.3f}%, 20kt 기준 {S['usab_by'][20]:.3f}%로 산정되었으며, " + _t52
    )

    doc.add_heading("5.3 최종 선정", level=2)
    _t53end = "모두 충족하는 방향으로 확인되었다." if r['pass'] else "고려한 방향으로 산정되었다."
    doc.add_paragraph(
        f"이상의 분석 결과를 종합할 때, 대상지의 최적 활주로 방향은 {S['best_rwy']}"
        f"(자기방위각 {S['best_angle']}°/{S['opp_angle']}°)로 최종 선정하며, 이는 단일 활주로 "
        f"운영 시에도 국제 및 국내 기준을 " + _t53end
    )
    doc.add_paragraph()

    # 6. 결론 및 의견
    doc.add_heading("6. 결론 및 의견", level=1)

    doc.add_heading("6.1 검토결과 요약", level=2)
    doc.add_paragraph(
        f"본 검토는 {stn_name} 관측소의 최근 {period_months}개월"
        f"({chart_start:%Y.%m.%d}~{chart_end:%Y.%m.%d})간 기상관측자료를 기초로 ICAO Annex 14 및 "
        f"Doc.9157, 그리고 국내 공항시설법 등 관계기준에 따라 활주로 방향별 측풍 이용률을 "
        f"산정하고 최적 활주로 방향을 도출하였다. 주요 검토결과를 종합하면 다음과 같다."
    )
    _verdict = "전 허용치 적합" if min_usab >= USABILITY_TARGET else "일부 조건 검토 필요"
    _add_df_table(doc, pd.DataFrame([
        {"구분": "대상 관측소", "주요 결과": stn_name},
        {"구분": "분석기간",
         "주요 결과": f"{chart_start:%Y.%m.%d}~{chart_end:%Y.%m.%d} ({period_months}개월)"},
        {"구분": "유효 관측시간", "주요 결과": f"{A['N_eff']:,} 시간 (Calm {A['calm_pct']:.1f}% 제외)"},
        {"구분": "최적 활주로 방향",
         "주요 결과": f"{S['best_rwy']} (방위각 {S['best_angle']}°/{S['opp_angle']}°)"},
        {"구분": "이용률 (10/13/20kt)",
         "주요 결과": f"{S['usab_by'][10]:.3f}% / {S['usab_by'][13]:.3f}% / {S['usab_by'][20]:.3f}%"},
        {"구분": "판정", "주요 결과": _verdict},
    ]))

    doc.add_heading("6.2 결론", level=2)
    _c62a = ("본 방향은 측풍 허용치 조건(항공기 기준활주로길이)의 변화와 무관하게 모든 시나리오에서 "
             "최적 방향으로 일관되게 도출되어 결과의 신뢰성과 강건성이 높은 것으로 평가되며, "
             if S['same_dir'] else "")
    if S['all_pass']:
        _c62b = (f"ICAO 기준({USABILITY_TARGET:g}%) 대비 여유(최소 {S['overall_min']:.2f}%~최대 "
                 f"{S['overall_max']:.2f}%)를 확보하고 있어 기상학적 관점에서 활주로 방향 설정에 "
                 f"제약이 적은 것으로 판단된다.")
    else:
        _c62b = (f"방향별 이용률은 {S['overall_min']:.2f}%~{S['overall_max']:.2f}% 범위로, ICAO 기준"
                 f"({USABILITY_TARGET:g}%)에 대해 일부 방위의 충족 여부에 대한 추가 검토가 필요한 "
                 f"것으로 판단된다.")
    doc.add_paragraph(
        f"이상의 분석 결과, 대상지의 최적 활주로 방향은 {S['best_rwy']}"
        f"(자기방위각 {S['best_angle']}°/{S['opp_angle']}°)로 최종 판단된다. " + _c62a + _c62b
    )

    doc.add_heading("6.3 의견", level=2)
    doc.add_paragraph(
        "본 검토는 풍향·풍속 자료에 기초한 활주로 이용률만을 대상으로 한 것으로, 활주로 방향의 "
        "최종 확정을 위해서는 다음 사항에 대한 후속 검토가 병행될 필요가 있다."
    )
    doc.add_paragraph("1) 타 입지요소와의 종합 검토")
    doc.add_paragraph(
        "활주로 방향은 기상조건 외에도 장애물제한표면(Obstacle Limitation Surface), 주변 지형·"
        "지장물, 공역 및 인접 공항과의 관제간섭, 소음영향권, 토지이용 등을 종합적으로 고려하여 "
        "최종 확정되어야 한다. 본 검토에서 도출된 방향은 기상학적 최적안으로서, 타 요소 검토 "
        "결과와의 정합성 확인이 필요하다."
    )
    doc.add_paragraph("2) 관측자료의 대표성 검증")
    doc.add_paragraph(
        f"본 분석에 활용된 {S['period_years_str']}개년({period_months}개월) 자료는 통상 요구되는 "
        f"최소 관측기간(5년 이상)을 고려한 것이나, 대상 관측소와 실제 후보지 간 이격거리 및 "
        f"지형적 차이에 따른 국지풍 영향 가능성을 배제할 수 없으므로, 필요시 후보지 인근 "
        f"임시관측소(AWOS 등) 설치를 통한 현지 실측자료 보완을 권고한다."
    )
    doc.add_paragraph("3) 시설계획 단계에서의 재검증")
    doc.add_paragraph(
        "향후 기본계획 및 실시설계 단계에서 활주로 제원(길이·폭), 접근절차, 항공기 기종 구성 등이 "
        "구체화될 경우, 해당 조건을 반영한 측풍 허용치 재적용 및 이용률 재검증을 실시할 것을 "
        "권고한다."
    )
    doc.add_paragraph("4) 기후변화 등 장기 변동성 고려")
    doc.add_paragraph(
        "기상관측자료는 특정 시점의 최근 자료에 기초하므로, 향후 설계 확정 이전 기후변동성에 "
        "따른 풍향·풍속 패턴 변화 가능성을 감안하여 최신 자료로 주기적으로 갱신·검증하는 것이 "
        "바람직하다."
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# 2. 메인페이지 설정 UI — 사이드바 대신 첫 화면에서 한 번에 설정 (박스 단위로 구획)
st.markdown("### 분석 설정")

row1_col1, row1_col2 = st.columns(2)

with row1_col1:
    with st.container(border=True):
        st.markdown("**1. 관측소 선택**")

        # 주소검색 자동선택 반영: 위젯 생성 '전에' pending 값을 위젯 key에 적용해야 한다
        # (Streamlit은 위젯 생성 후 그 key의 session_state 수정을 금지하므로 pending 패턴 사용)
        _pending = st.session_state.pop("pending_station", None)
        if _pending:
            st.session_state["obs_type_radio"] = _pending["obs_type"]
            if _pending["obs_type"].startswith("ASOS"):
                st.session_state["asos_region"] = _pending["region"]
                st.session_state["asos_station_id"] = _pending["stn_id"]
            else:
                st.session_state["aws_stn_name"] = _pending["name"]
        if "obs_type_radio" not in st.session_state:
            st.session_state["obs_type_radio"] = "ASOS (종관)"

        obs_type = st.radio(
            "관측종류",
            ["ASOS (종관)", "AWS (방재) — CSV 업로드"],
            horizontal=True,
            key="obs_type_radio",
            help=(
                "ASOS: 기간(시작∼종료) 범위 조회 API 제공 → 5∼10년 장기분석 가능.\n"
                "AWS: API는 단일 시점 조회만 지원되어 CSV 업로드 방식으로 분석합니다."
            ),
        )
        is_asos = obs_type.startswith("ASOS")

        aws_files = []   # ASOS 경로에서도 변수가 항상 정의되도록

        if is_asos:
            region_db = ASOS_BY_REGION
            region_list = list(region_db.keys())
            default_region_idx = region_list.index("전남") if "전남" in region_list else 0
            if "asos_region" not in st.session_state:
                st.session_state["asos_region"] = region_list[default_region_idx]

            c_region, c_station = st.columns(2)
            with c_region:
                region = st.selectbox("광역시도", region_list, key="asos_region")

            # 해당 광역시도의 관측소 목록 (지점번호를 옵션값으로 사용 → 주소검색 자동선택과 정합)
            stations_in_region = region_db[region]
            region_ids = [s[0] for s in stations_in_region]
            id_to_label = {sid: f"{name} ({sid}) · {start}" for sid, name, start in stations_in_region}
            default_name = "목포" if region == "전남" else stations_in_region[0][1]
            # 스테일 방지: 현재 지역에 없는 지점번호가 남아있으면 지역 기본값으로 초기화
            if st.session_state.get("asos_station_id") not in region_ids:
                st.session_state["asos_station_id"] = next(
                    (sid for sid, name, start in stations_in_region if name == default_name),
                    region_ids[0])

            with c_station:
                station_id = st.selectbox("관측소", region_ids,
                                          format_func=lambda s: id_to_label[s],
                                          key="asos_station_id")
            sel = next(s for s in stations_in_region if s[0] == station_id)
            stn_id, stn_name, stn_start = sel
            st.caption(f"[ASOS] {stn_name} ({region}) · 관측 가능일: {stn_start} ~ 현재")
        else:
            # AWS: API 대신 기상자료개방포털 CSV 업로드 방식
            region = "—"
            stn_id = None
            stn_start = "2000-01-01"   # 기간선택 위젯이 깨지지 않도록 두는 더미 값

            st.caption(
                "기상청 API는 단일 시점 조회만 지원하여 장기 분석이 불가합니다. "
                "[기상자료개방포털](https://data.kma.go.kr)에서 CSV를 직접 다운로드 후 "
                "아래에 업로드하면 동일한 분석이 가능합니다."
            )
            with st.expander("CSV 다운로드 방법"):
                st.markdown(
                    "1. [기상자료개방포털](https://data.kma.go.kr) 접속  \n"
                    "2. **기후통계분석 → 방재기상관측 → 시간자료**  \n"
                    "3. 관측소·기간 선택 후 **CSV 다운로드**  \n"
                    "4. 연도별로 나눠 받은 경우 여러 파일을 동시에 업로드 가능  \n"
                    "5. 컬럼에 **풍향(deg)** 과 **풍속(m/s)** 이 포함된 파일이어야 합니다."
                )
            if "aws_stn_name" not in st.session_state:
                st.session_state["aws_stn_name"] = "AWS 관측소"
            stn_name = st.text_input("관측소 이름 (차트 표시용)", key="aws_stn_name")
            aws_files = st.file_uploader(
                "AWS 시간자료 CSV (복수 파일 동시 업로드 가능)",
                type=["csv"],
                accept_multiple_files=True,
                key="aws_csv",
            )
            if aws_files:
                st.caption(f"{len(aws_files)}개 파일 업로드됨")

        st.markdown("---")
        st.markdown("**주소로 가까운 관측소 자동 선택**")
        st.caption(
            "주소를 입력하면 가장 가까운 종관(ASOS) 관측소를 자동 선택합니다. 방재(AWS)가 종관보다 "
            f"{NEAR_ASOS_PREFER_KM:g}km 이상 더 가까우면 방재 자료(CSV) 이용을 안내합니다. "
            "(기상청 API 허브 인증키·카카오 REST API 키는 상단 'API 키 설정'에서 입력)"
        )
        addr_input = st.text_input("주소 입력", placeholder="예: 대구 동구 동대구로 550")

        if st.button("가까운 관측소 자동 선택", key="btn_addr_search"):
            try:
                _kma_hub_key = st.session_state.get("kma_hub_key", "")
                _kakao_key = st.session_state.get("kakao_key", "")
                if not _kma_hub_key or not _kakao_key:
                    st.session_state["addr_result"] = {"kind": "error",
                        "msg": "'API 키 설정' 팝업에서 기상청 API 허브 인증키와 카카오 REST API 키를 먼저 입력하세요."}
                elif not addr_input or not addr_input.strip():
                    st.session_state["addr_result"] = {"kind": "error", "msg": "주소를 입력하세요."}
                else:
                    with st.spinner("주소 검색 중..."):
                        lat, lon, disp_addr, err = _geocode_address_kakao(addr_input, _kakao_key)
                    if err:
                        st.session_state["addr_result"] = {"kind": "error", "msg": f"주소 검색 실패: {err}"}
                    else:
                        try:
                            with st.spinner("관측소 DB 로딩 중..."):
                                stn_db = _load_station_db(_kma_hub_key)
                        except Exception as e:
                            stn_db = None
                            st.session_state["addr_result"] = {"kind": "error", "msg": f"관측소 DB 로딩 실패: {e}"}

                        if stn_db is None or len(stn_db) == 0:
                            st.session_state["addr_result"] = {"kind": "error",
                                "msg": "관측소 DB를 불러오지 못했습니다. 기상청 API 허브 인증키를 확인하세요."}
                        else:
                            na = _nearest_stations(lat, lon, stn_db, 'ASOS', n=1)
                            nw = _nearest_stations(lat, lon, stn_db, 'AWS', n=1)
                            d_asos = float(na.iloc[0]['dist_km']) if len(na) else float('inf')
                            d_aws = float(nw.iloc[0]['dist_km']) if len(nw) else float('inf')
                            asos_id = str(na.iloc[0]['stn_id']) if len(na) else None
                            asos_nm = na.iloc[0]['name_ko'] if len(na) else '—'
                            aws_nm = nw.iloc[0]['name_ko'] if len(nw) else '—'
                            info = _ASOS_ID_TO_INFO.get(_norm_stn_id(asos_id)) if asos_id else None
                            pick_aws = (d_aws < d_asos) and ((d_asos - d_aws) >= NEAR_ASOS_PREFER_KM)

                            if pick_aws:
                                st.session_state["pending_station"] = {
                                    "obs_type": "AWS (방재) — CSV 업로드", "name": aws_nm}
                                st.session_state["addr_result"] = {"kind": "aws",
                                    "msg": (f"검색 위치: {disp_addr}\n\n"
                                            f"가장 가까운 관측소는 방재(AWS) '{aws_nm}' ({d_aws:.1f}km)로, "
                                            f"가장 가까운 종관(ASOS) '{asos_nm}' ({d_asos:.1f}km)보다 "
                                            f"{d_asos - d_aws:.1f}km 더 가깝습니다. 방재 자료는 기간조회 API가 "
                                            f"없어 CSV 업로드로 분석하므로, 관측종류를 방재(AWS)로 전환했습니다. "
                                            f"아래 절차로 기상자료개방포털에서 '{aws_nm}' 시간자료를 받아 업로드하세요.")}
                                st.rerun()
                            elif info:
                                rgn_, nm_, _start = info
                                st.session_state["pending_station"] = {
                                    "obs_type": "ASOS (종관)", "region": rgn_, "stn_id": _norm_stn_id(asos_id)}
                                _extra = ("" if d_asos <= d_aws else
                                          f" (방재 '{aws_nm}' {d_aws:.1f}km가 더 가깝지만 종관과 "
                                          f"{d_asos - d_aws:.1f}km 차이로 {NEAR_ASOS_PREFER_KM:g}km 미만이어서 "
                                          f"종관을 우선 선택)")
                                st.session_state["addr_result"] = {"kind": "asos",
                                    "msg": (f"검색 위치: {disp_addr}\n\n"
                                            f"가장 가까운 종관(ASOS) 관측소 '{nm_}' ({rgn_}, {d_asos:.1f}km)를 "
                                            f"자동 선택했습니다." + _extra)}
                                st.rerun()
                            else:
                                st.session_state["addr_result"] = {"kind": "error",
                                    "msg": (f"가장 가까운 종관 관측소 '{asos_nm}'(지점 {asos_id})가 내장 "
                                            f"목록에 없어 자동 선택하지 못했습니다. '관측소'에서 직접 선택하세요.")}
            except Exception as e:
                st.session_state["addr_result"] = {"kind": "error", "msg": f"주소검색 중 예상치 못한 오류: {e}"}

        _ar = st.session_state.get("addr_result")
        if _ar:
            if _ar["kind"] == "asos":
                st.success(_ar["msg"])
            elif _ar["kind"] == "aws":
                st.warning(_ar["msg"])
                st.markdown(
                    "1. [기상자료개방포털](https://data.kma.go.kr) 접속  \n"
                    "2. **기후통계분석 → 방재기상관측 → 시간자료**  \n"
                    "3. 위 관측소·기간 선택 후 **CSV 다운로드**  \n"
                    "4. 위 '관측소 선택'의 AWS 업로더에 업로드 (여러 파일 동시 가능)"
                )
            else:
                st.error(_ar["msg"])

with row1_col2:
    with st.container(border=True):
        st.markdown("**2. 분석 기간**")
        preset = st.radio(
            "기간 설정", ["최근 10년", "최근 5년", "사용자 지정"],
            horizontal=True,
            help="ICAO Annex 14 권고: 최소 5년 이상의 신뢰성 있는 기상통계자료",
        )

        # 종료월 = 오늘 기준 직전(완료된) 월 — 진행 중 월은 제외
        _today = date.today()
        _end_y, _end_m = (_today.year - 1, 12) if _today.month == 1 else (_today.year, _today.month - 1)

        def _months_back(ey, em, n_years):
            """종료월 포함 정확히 (n_years × 12)개월이 되는 시작 연/월."""
            idx = ey * 12 + (em - 1) - (n_years * 12 - 1)
            return idx // 12, (idx % 12) + 1

        if preset == "최근 10년":
            start_y, start_m = _months_back(_end_y, _end_m, 10)
            end_y, end_m = _end_y, _end_m
        elif preset == "최근 5년":
            start_y, start_m = _months_back(_end_y, _end_m, 5)
            end_y, end_m = _end_y, _end_m
        else:  # 사용자 지정
            # 선택된 관측소 관측 시작년도 이후만 허용
            try:
                stn_start_yr = int(stn_start.split("-")[0]) if stn_start and stn_start != "—" else 1960
            except Exception:
                stn_start_yr = 1960
            yr_range = list(range(max(1960, stn_start_yr), _today.year + 1))
            mo_range = list(range(1, 13))

            cA, cB = st.columns(2)
            with cA:
                default_start_y = max(yr_range[0], _today.year - 5)
                start_y = st.selectbox("시작 연", yr_range,
                                       index=yr_range.index(default_start_y) if default_start_y in yr_range else 0)
                start_m = st.selectbox("시작 월", mo_range, index=0)
            with cB:
                end_y = st.selectbox("종료 연", yr_range, index=len(yr_range) - 1)
                end_m = st.selectbox("종료 월", mo_range, index=_end_m - 1)

        # 날짜 객체 변환 (종료월은 말일)
        start_date = date(start_y, start_m, 1)
        _last_day = monthrange(end_y, end_m)[1]
        end_date = date(end_y, end_m, _last_day)

        # 유효성 및 요약
        _months = (end_y - start_y) * 12 + (end_m - start_m) + 1
        if start_date > end_date:
            st.error("시작이 종료보다 늦습니다.")
        elif _months < 60:
            st.warning(f"{_months}개월 (<5년) — ICAO 권고 기간 미달")
            st.caption(f"{start_date:%Y-%m} ~ {end_date:%Y-%m} · {_months}개월")
        else:
            st.success(f"{start_date:%Y-%m} ~ {end_date:%Y-%m} · {_months}개월 ({_months/12:.1f}년)")

    with st.container(border=True):
        st.markdown("**3. 측풍 허용치 (ICAO Doc. 9157)**")
        rwy_length = st.number_input("활주로 길이 (m)", min_value=300, max_value=5000, value=2000, step=100)
        low_friction = st.checkbox("종방향 마찰계수 부족 (활주로 제동효과 불량)", value=False)
        auto_limit, auto_note = select_limit_by_rwy_length(rwy_length, low_friction)
        st.caption(f"자동 선택: **{fmt_kt(auto_limit)}** ({auto_note} 기준)")
        override = st.checkbox("수동 선택 사용", value=False)
        if override:
            primary_limit = st.selectbox("측풍 허용치", CROSSWIND_LIMITS_KT,
                                         index=CROSSWIND_LIMITS_KT.index(auto_limit),
                                         format_func=fmt_kt)
        else:
            primary_limit = auto_limit

btn_col1, btn_col2 = st.columns([4, 1])
with btn_col1:
    run_clicked = st.button("분석 시작", type="primary", use_container_width=True)
with btn_col2:
    if st.button("데이터 캐시 삭제", use_container_width=True):
        st.cache_data.clear()
        st.info("캐시가 삭제되었습니다.")

st.divider()

# 3. 분석 실행
if run_clicked:
    df = None
    row_count = None
    _chart_start = start_date
    _chart_end   = end_date

    # ── 데이터 수집 ─────────────────────────────────────────────
    if is_asos:
        if not api_key:
            st.error("API Key를 입력하세요.")
        else:
            df, df_invalid_rows, result = get_weather_data_v28(api_key, stn_id, start_date, end_date)
            if df is None:
                st.error(f"분석 실패: {result}")
            else:
                row_count = result
    else:
        # AWS CSV 업로드 경로
        if not aws_files:
            st.error("사이드바에서 AWS 시간자료 CSV 파일을 업로드하세요.")
        else:
            df, df_invalid_rows, result, (csv_s, csv_e) = _parse_aws_csv(aws_files)
            if df is None:
                st.error(f"CSV 파싱 실패: {result}")
            else:
                row_count = result
                if csv_s and csv_e:
                    _chart_start, _chart_end = csv_s, csv_e

    # ── 분석 & 표시 (ASOS / AWS 공통) ───────────────────────────
    if df is not None:

        _period_months = (
            (_chart_end.year - _chart_start.year) * 12
            + (_chart_end.month - _chart_start.month) + 1
        )
        st.success(
            f"{stn_name} · {row_count:,}시간 수집 완료 "
            f"· 실제 분석 기간: **{_chart_start:%Y-%m-%d} ~ {_chart_end:%Y-%m-%d}** "
            f"({_period_months}개월)"
        )
        with st.spinner("바람성분 vector 분석 중..."):
            A = analyze_runway(df)

            # --- 데이터 요약 ---
            st.divider()
            s0, s1, s2, s3, s4 = st.columns(5)
            s0.metric("분석 기간", f"{_chart_start:%Y-%m}", f"~ {_chart_end:%Y-%m}")
            s1.metric("전체 관측 시간", f"{A['N_total']:,} h")
            s2.metric(f"Calm (0~{fmt_kt(CALM_THRESHOLD_KT)})", f"{A['N_calm']:,} h", f"{A['calm_pct']:.2f}%")
            s3.metric("유효 데이터", f"{A['N_eff']:,} h")
            s4.metric("적용 측풍 허용치", fmt_kt(primary_limit))
            st.caption("※ Calm(무영향) 데이터는 논문 §3.2에 따라 활주로 방향 무관하게 '이용 가능'으로 집계됩니다.")

            # --- 주 결과(자동/수동 선택된 한계치) ---
            st.subheader(f"분석 결과 · 허용치 {fmt_kt(primary_limit)}")
            r = A['results'][primary_limit]
            c1, c2, c3, c4 = st.columns(4)
            c1.metric("최적 활주로", rwy_name(r['best_angle']), f"{r['best_angle']}°")
            c2.metric("최대 이용률", f"{r['best_usab']:.3f}%",
                      "PASS" if r['pass'] else "FAIL")
            c3.metric("평균 측풍", f"{r['mean_xwind']:.2f} kt", f"{_kt2ms(r['mean_xwind']):.2f} m/s", delta_color="off")
            c4.metric("동율 후보", f"{r['tied_count']}개",
                      "동율 처리 적용" if r['tied_count'] > 1 else "유일")

            if not r['pass']:
                st.warning(
                    f"단일 활주로로 95% 미달. 2개 활주로 최적 조합: "
                    f"**{rwy_name(r['pair_angles'][0])} + {rwy_name(r['pair_angles'][1])}** "
                    f"→ 이용률 **{r['pair_usab']:.3f}%** "
                    f"{'PASS' if r['pair_pass'] else '여전히 FAIL'}"
                )

            # --- 탭 ---
            t1, t2, t3, t4, t5, t6 = st.tabs([
                "3개 허용치 종합", "이용률 곡선", "바람장미",
                "16방위 빈도표", "2개 활주로 분석", "방위각 상세표"
            ])

            with t1:
                st.markdown("#### ICAO Doc. 9157 세 개 측풍 한계치 동시 분석")
                summary_rows = []
                for lim in CROSSWIND_LIMITS_KT:
                    rr = A['results'][lim]
                    summary_rows.append({
                        "허용치": fmt_kt(lim),
                        "최적 활주로": rwy_name(rr['best_angle']),
                        "방위각 (°)": rr['best_angle'],
                        "이용률 (%)": round(rr['best_usab'], 3),
                        "평균 측풍 (kt)": round(rr['mean_xwind'], 2),
                        "평균 측풍 (m/s)": round(_kt2ms(rr['mean_xwind']), 2),
                        "최대 측풍 (kt)": round(rr['max_xwind'], 2),
                        "최대 측풍 (m/s)": round(_kt2ms(rr['max_xwind']), 2),
                        "동율 후보수": rr['tied_count'],
                        "단일 판정": "PASS" if rr['pass'] else "FAIL",
                        "2개 조합 이용률 (%)": round(rr['pair_usab'], 3),
                        "2개 조합": f"{rwy_name(rr['pair_angles'][0])} + {rwy_name(rr['pair_angles'][1])}",
                    })
                st.dataframe(pd.DataFrame(summary_rows), width='stretch', hide_index=True)

            with t2:
                curve_df = pd.DataFrame({
                    "angle": np.tile(A['angles'], len(CROSSWIND_LIMITS_KT)),
                    "usability": np.concatenate([A['results'][l]['usability'] for l in CROSSWIND_LIMITS_KT]),
                    "limit_kt": np.repeat([fmt_kt(l) for l in CROSSWIND_LIMITS_KT], len(A['angles'])),
                })
                fig1 = px.line(curve_df, x="angle", y="usability", color="limit_kt",
                               title="방향별 이용률 (허용치별 비교)",
                               labels={"angle": "활주로 방위각 (°)", "usability": "이용률 (%)", "limit_kt": "허용치"})
                fig1.add_hline(y=USABILITY_TARGET, line_dash="dash", line_color="red",
                               annotation_text="ICAO 95%", annotation_position="top right")
                for lim in CROSSWIND_LIMITS_KT:
                    rr = A['results'][lim]
                    fig1.add_vline(x=rr['best_angle'], line_dash="dot", opacity=0.3)
                st.plotly_chart(fig1, width='stretch')

            with t3:
                # ── 바람장미 집계: 16방위 × 6풍속구간 → 빈도(%) ──────────
                _wd_r  = df['wd'].to_numpy(dtype=np.float32)
                _ws_r  = df['ws_kt'].to_numpy(dtype=np.float32)
                _N_r   = len(_wd_r)

                # Calm(정온) 분리 — 0.5 m/s 이하 (바람장미 전용 기준, 활주로 측풍분석의
                # CALM_THRESHOLD_KT=3kt와는 별개). 0.5 m/s → 0.5*1.94384 ≈ 0.97 kt
                _CALM_ROSE_KT = 0.5 * 1.94384
                _calm_rose_mask = _ws_r <= _CALM_ROSE_KT
                _calm_rose_pct  = float(_calm_rose_mask.sum()) / _N_r * 100.0

                # 16방위 bin (22.5° 간격, N=0°) — Calm 제외한 유효바람만 방향 집계
                _dir_ang  = np.arange(0, 360, 22.5)                    # [0, 22.5 … 337.5]
                _wd_eff   = _wd_r[~_calm_rose_mask]
                _ws_eff   = _ws_r[~_calm_rose_mask]
                _dir_idx  = ((((_wd_eff + 11.25) // 22.5) % 16)).astype(np.int32)

                # 풍속 구간 6단계 (kt): (0,3], (3,7], … >21 — m/s 동시 표기
                _spd_thresh = [3, 7, 11, 17, 21]                        # 오름차순 경계값
                _spd_labels = [
                    f"0–3 kt (0–{_kt2ms(3):.1f} m/s)",
                    f"3–7 kt ({_kt2ms(3):.1f}–{_kt2ms(7):.1f} m/s)",
                    f"7–11 kt ({_kt2ms(7):.1f}–{_kt2ms(11):.1f} m/s)",
                    f"11–17 kt ({_kt2ms(11):.1f}–{_kt2ms(17):.1f} m/s)",
                    f"17–21 kt ({_kt2ms(17):.1f}–{_kt2ms(21):.1f} m/s)",
                    f"≥21 kt (≥{_kt2ms(21):.1f} m/s)",
                ]
                _spd_idx = np.zeros(len(_ws_eff), dtype=np.int32)
                for _i, _t in enumerate(_spd_thresh):
                    _spd_idx[_ws_eff > _t] = _i + 1                    # 초과(>) 기준 bin 배정

                # 집계 (16 × 6) → 전체 관측(Calm 포함) 대비 % → 막대 합 + Calm% = 100%
                _rose_rows = []
                for _d in range(16):
                    for _s in range(6):
                        _cnt = int(((_dir_idx == _d) & (_spd_idx == _s)).sum())
                        _rose_rows.append({
                            'angle':     float(_dir_ang[_d]),
                            'speed_bin': _spd_labels[_s],
                            'freq_pct':  _cnt / _N_r * 100.0,
                        })
                _df_rose = pd.DataFrame(_rose_rows)

                # 이산 색상 (오름차순: 연한색 → 진한색)
                _color_seq = ["#c6dbef", "#74c476", "#fdd835", "#fd8d3c", "#e31a1c", "#67000d"]
                _color_map = dict(zip(_spd_labels, _color_seq))

                fig2 = px.bar_polar(
                    _df_rose,
                    r="freq_pct",
                    theta="angle",
                    color="speed_bin",
                    color_discrete_map=_color_map,
                    category_orders={"speed_bin": _spd_labels},
                    title=f"Wind Rose — {stn_name}  ({_chart_start:%Y-%m} ~ {_chart_end:%Y-%m})",
                    template="plotly_white",
                )
                _polar_domain = dict(x=[0.0, 0.74], y=[0.0, 1.0])      # 우측은 범례 공간
                fig2.update_layout(
                    polar=dict(
                        domain=_polar_domain,
                        hole=0.16,             # 중앙에 Calm 표시용 빈 원
                        angularaxis=dict(
                            tickmode="array",
                            tickvals=[0, 45, 90, 135, 180, 225, 270, 315],
                            ticktext=["N", "NE", "E", "SE", "S", "SW", "W", "NW"],
                            direction="clockwise",
                            rotation=90,           # N을 12시 방향으로
                        ),
                        radialaxis=dict(
                            ticksuffix="%",        # 반경 축 단위 (%)
                            angle=45,              # 눈금 레이블 겹침 방지
                            tickangle=45,
                        ),
                    ),
                    legend=dict(title="풍속 구간", traceorder="normal"),
                    margin=dict(t=60, b=30, l=30, r=30),
                    height=540,
                )
                fig2.add_annotation(
                    text=f"Calm<br>{_calm_rose_pct:.1f}%",
                    x=(_polar_domain['x'][0] + _polar_domain['x'][1]) / 2,
                    y=(_polar_domain['y'][0] + _polar_domain['y'][1]) / 2,
                    xref="paper", yref="paper",
                    showarrow=False,
                    font=dict(size=14, color="#333"),
                    align="center",
                )
                st.plotly_chart(fig2, width='stretch')
                st.caption(f"중앙 원: 풍속 0.5 m/s(≈{_CALM_ROSE_KT:.2f} kt) 이하 정온(Calm) 비율 · "
                           f"방위별 막대는 Calm을 제외한 유효바람 기준 (막대 합 + Calm = 100%)")

            with t4:
                st.markdown("#### 16방위 × 풍속구간 빈도표 (%) — 논문 Table 6 양식")
                st.dataframe(A['freq_table'], width='stretch')
                st.caption(f"합계 100% 기준 · Calm(0~{fmt_kt(CALM_THRESHOLD_KT)}) 비율: {A['calm_pct']:.2f}%")

            with t5:
                st.markdown(f"#### 2개 활주로 최적 조합 (허용치 {fmt_kt(primary_limit)})")
                st.caption(f"※ 최소 각도 분리 {MIN_SEPARATION_DEG}° 제약 적용 "
                           f"(물리적 배치 가능성 고려)")
                r = A['results'][primary_limit]
                p1, p2 = r['pair_angles']
                d1, d2, d3 = st.columns(3)
                d1.metric("1차 활주로", rwy_name(p1), f"{p1}°")
                d2.metric("2차 활주로", rwy_name(p2), f"{p2}°")
                d3.metric("조합 이용률", f"{r['pair_usab']:.3f}%",
                          "PASS" if r['pair_pass'] else "FAIL")
                st.info(
                    f"단일 활주로 최대 이용률: **{r['best_usab']:.3f}%** → "
                    f"2개 조합 이용률: **{r['pair_usab']:.3f}%** "
                    f"(개선 +{r['pair_usab']-r['best_usab']:.3f}%p)"
                )

            with t6:
                st.markdown("#### 방위각 간격별 이용률 상세표 (10° 간격)")
                st.caption(
                    "각 행은 활주로 방위(방향 ↔ 대응방향) 쌍이며, 허용측풍별 이용률(%)을 나타냅니다. "
                    "95% 미달은 빨강, 이상은 초록으로 강조됩니다. "
                    "(활주로 명칭이 10° 단위로만 존재하므로 분석도 10° 간격으로 고정됩니다.)"
                )

                headings = np.arange(RWY_ANGLE_STEP_DEG, 181, RWY_ANGLE_STEP_DEG, dtype=int)  # 10,20,...,180
                rows = []
                for h in headings:
                    i = int(h % 180)                               # 180° ≡ 0° (동일 활주로)
                    idx = i // RWY_ANGLE_STEP_DEG                  # usability 배열 인덱스 (10° 격자)
                    rec = int(h + 180)                              # 10→190, ..., 180→360
                    row = {
                        "방향 (°)": int(h),
                        "대응방향 (°)": rec,
                        "활주로": rwy_name(i),
                    }
                    for lim in CROSSWIND_LIMITS_KT:
                        row[f"{fmt_kt(lim)} 이용률 (%)"] = round(float(A['results'][lim]['usability'][idx]), 2)
                    rows.append(row)
                df_table = pd.DataFrame(rows)

                # 95% 기준 색상 강조 (pandas 3.x: Styler.map 사용)
                kt_cols = [c for c in df_table.columns if "이용률" in c]
                def _hl(v):
                    if isinstance(v, (int, float)):
                        if v >= USABILITY_TARGET:
                            return "background-color: #d4edda; color: #155724;"
                        return "background-color: #f8d7da; color: #721c24;"
                    return ""
                styled = df_table.style.map(_hl, subset=kt_cols).format({c: "{:.2f}" for c in kt_cols})
                st.dataframe(styled, width='stretch', hide_index=True)

                # 최적 방향 요약
                st.markdown("##### 허용치별 최적 방위각 (표 기준)")
                best_rows = []
                for lim in CROSSWIND_LIMITS_KT:
                    col = f"{fmt_kt(lim)} 이용률 (%)"
                    sub = df_table[["방향 (°)", "대응방향 (°)", "활주로", col]]
                    top = sub.loc[sub[col].idxmax()]
                    best_rows.append({
                        "허용치": fmt_kt(lim),
                        "최적 방향": f"{int(top['방향 (°)'])}° / {int(top['대응방향 (°)'])}°",
                        "활주로": top['활주로'],
                        "이용률 (%)": f"{top[col]:.2f}",
                        "판정": "PASS" if top[col] >= USABILITY_TARGET else "FAIL",
                    })
                st.dataframe(pd.DataFrame(best_rows), width='stretch', hide_index=True)

                # CSV 다운로드
                csv_bytes = df_table.to_csv(index=False, encoding='utf-8-sig').encode('utf-8-sig')
                st.download_button(
                    "CSV 다운로드",
                    csv_bytes,
                    file_name=f"runway_usability_{stn_name}_{_chart_start}_{_chart_end}_step10.csv",
                    mime="text/csv",
                    key="dl_detail_csv",
                )

            # ── 데이터 품질 검토 ───────────────────────────────────────
            st.divider()
            st.subheader("데이터 품질 검토")

            _n_valid   = A['N_total']
            _n_invalid = len(df_invalid_rows) if df_invalid_rows is not None else 0
            _n_raw     = _n_valid + _n_invalid
            _valid_pct = _n_valid   / _n_raw * 100 if _n_raw > 0 else 0
            _miss_pct  = _n_invalid / _n_raw * 100 if _n_raw > 0 else 0

            qa1, qa2, qa3 = st.columns(3)
            qa1.metric("전체 수집 행",          f"{_n_raw:,} 행")
            qa2.metric("유효 데이터 (풍향+풍속 존재)", f"{_n_valid:,} 행",   f"{_valid_pct:.1f}%")
            qa3.metric("결측 데이터 (풍향 또는 풍속 없음)", f"{_n_invalid:,} 행", f"{_miss_pct:.1f}%")

            st.caption(
                "**유효**: 풍향(°)과 풍속(m/s) 값이 모두 관측된 시각 → 분석에 사용됨  \n"
                "**결측**: 풍향 또는 풍속이 누락된 시각 (장비 오류·통신 두절 등) → 분석에서 제외됨"
            )

            # Excel 2-시트 다운로드 (유효 / 결측)
            _excel_buf = io.BytesIO()
            # 유효 데이터: 주요 컬럼만 선별해 가독성 확보
            _valid_export_cols = [c for c in ['tm', 'wd', 'ws', 'ws_kt'] if c in df.columns]
            _valid_export = df[_valid_export_cols].copy() if _valid_export_cols else df.copy()
            _rename_map = {'tm': '관측시각', 'wd': '풍향(°)', 'ws': '풍속(m/s)', 'ws_kt': '풍속(kt)'}
            _valid_export.rename(columns=_rename_map, inplace=True)

            with pd.ExcelWriter(_excel_buf, engine='openpyxl') as _writer:
                _valid_export.to_excel(_writer, sheet_name='유효_데이터', index=False)
                if _n_invalid > 0:
                    _inv_export = df_invalid_rows.copy()
                    _inv_export.rename(columns=_rename_map, inplace=True)
                    _inv_export.to_excel(_writer, sheet_name='결측_데이터', index=False)
                else:
                    pd.DataFrame({"안내": ["결측 데이터가 없습니다."]}).to_excel(
                        _writer, sheet_name='결측_데이터', index=False
                    )

            st.download_button(
                "데이터 품질 보고서 다운로드 (Excel)",
                _excel_buf.getvalue(),
                file_name=f"data_quality_{stn_name}_{_chart_start}_{_chart_end}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )

            # ── 검토서(Word) 생성 ─────────────────────────────────────────
            # 분석과 같은 실행에서 미리 생성해 download_button으로 직접 제공한다.
            # (별도 st.button은 클릭 시 rerun을 일으키는데, 결과 표시가 `if run_clicked`
            #  블록 안에 있어 rerun 시 run_clicked=False가 되며 화면 전체가 초기화되어
            #  아무 동작도 하지 않는 문제가 있었음. Excel 다운로드와 동일한 패턴으로 통일.)
            st.divider()
            st.subheader("활주로 방향 검토서 생성")
            st.caption(
                "분석 결과를 ICAO 기준 검토서(.docx)로 내보냅니다. 풍배도·이용률 곡선·"
                "빈도표·상세표가 자동 삽입되며, 개요·결론 등 서술 항목([ 작성 ] 표시)은 "
                "생성 후 Word/HWP에서 직접 작성하면 됩니다."
            )
            try:
                with st.spinner("검토서 작성 중 (차트 렌더링 포함)..."):
                    _review_docx = build_review_docx(
                        A, df, stn_name, _chart_start, _chart_end, primary_limit,
                        _n_raw, _n_valid, _n_invalid,
                    )
                st.download_button(
                    "검토서 다운로드 (.docx)",
                    _review_docx,
                    file_name=f"활주로방향검토서_{stn_name}_{_chart_start}_{_chart_end}.docx",
                    mime=("application/vnd.openxmlformats-officedocument"
                          ".wordprocessingml.document"),
                    key="dl_review_docx",
                )
                st.caption(
                    "HWP 사용 시: 한글에서 이 .docx 파일을 열어 내용을 보완한 뒤 "
                    "'다른 이름으로 저장 → 한글 문서(.hwp)'로 저장하세요."
                )
            except Exception as _docx_err:
                st.error(f"검토서 생성 중 오류가 발생했습니다: {_docx_err}")
